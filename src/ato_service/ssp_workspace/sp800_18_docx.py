"""NIST SP 800-18 Revision 2 Table 1 shaped DOCX rendering for approved SSP snapshots."""

from __future__ import annotations

import json
import re
from typing import Any

from docx.document import Document as DocxDocument

SP800_18_SOURCE_LABEL = "NIST SP 800-18 Revision 2 (Table 1)"


def nist_control_sort_key(control_id: str) -> tuple[str | int, ...]:
    """Sort control IDs in catalog order (e.g. AC-3 before AC-11)."""
    parts = control_id.split("-", 1)
    if len(parts) != 2:
        return (control_id,)
    family, remainder = parts
    tokens: list[str | int] = [family]
    for segment in re.split(r"[().]", remainder):
        if not segment:
            continue
        if segment.isdigit():
            tokens.append(int(segment))
        else:
            tokens.append(segment)
    return tuple(tokens)


def render_sp800_18_docx(document: DocxDocument, snapshot: dict[str, Any]) -> None:
    """Append SP 800-18 Table 1 chapters to an empty document."""
    coverage = snapshot.get("standard_coverage") or []
    if not coverage:
        raise ValueError("standard_coverage is required for SP 800-18 DOCX rendering")

    controls_by_id = {item["control_id"]: item for item in snapshot["controls"]}
    ordered_controls = _ordered_controls(snapshot, controls_by_id)
    item_values = _item_value_lookup(snapshot)

    for entry in coverage:
        requirement_id = entry["requirement_id"]
        title = entry["title"]
        kind = entry["coverage_kind"]
        document.add_heading(title, level=1)

        if kind == "controls":
            if requirement_id == "table1.control-implementation-status":
                _render_control_status_table(document, ordered_controls, snapshot)
                continue
            if requirement_id == "table1.control-implementation-details":
                _render_control_details_chapter(
                    document,
                    snapshot,
                    item_values,
                    ordered_controls,
                )
                continue
            document.add_paragraph("Not provided.")
            continue

        if requirement_id == "table1.system-categorization":
            _render_fips199_table(document, item_values, snapshot)
            continue

        item_ids = entry.get("item_ids") or []
        for item_id in item_ids:
            label = _item_title(snapshot, item_id)
            document.add_heading(label, level=2)
            document.add_paragraph(_format_item_text(item_values.get(item_id)))


def _ordered_controls(
    snapshot: dict[str, Any],
    controls_by_id: dict[str, dict[str, Any]],
) -> list[dict[str, Any]]:
    order = snapshot.get("control_order") or []
    if order:
        ordered = [controls_by_id[cid] for cid in order if cid in controls_by_id]
        seen = {item["control_id"] for item in ordered}
        extras = [
            controls_by_id[cid]
            for cid in sorted(controls_by_id, key=nist_control_sort_key)
            if cid not in seen
        ]
        return ordered + extras
    return sorted(controls_by_id.values(), key=lambda item: nist_control_sort_key(item["control_id"]))


def _item_value_lookup(snapshot: dict[str, Any]) -> dict[str, Any]:
    values: dict[str, Any] = dict(snapshot.get("facts") or {})
    for section in snapshot.get("sections") or []:
        section_id = section["section_id"]
        if section_id not in values and section.get("content"):
            values[section_id] = section["content"]
    return values


def _item_title(snapshot: dict[str, Any], item_id: str) -> str:
    for item in snapshot.get("ssp_items") or []:
        if item.get("item_id") == item_id:
            return str(item.get("title") or item_id)
    return item_id.replace("_", " ").replace(".", " ").title()


def _format_item_text(value: Any) -> str:
    if value is None:
        return "Not provided."
    if isinstance(value, str):
        stripped = value.strip()
        return stripped if stripped else "Not provided."
    if isinstance(value, list):
        if not value:
            return "Not provided."
        lines: list[str] = []
        for entry in value:
            if isinstance(entry, str):
                lines.append(f"- {entry.strip()}" if entry.strip() else "- (empty)")
            elif isinstance(entry, dict):
                lines.append(f"- {_format_item_text(entry)}")
            else:
                lines.append(f"- {entry}")
        return "\n".join(lines)
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False, indent=2)
    return str(value)


def _render_fips199_table(
    document: DocxDocument,
    item_values: dict[str, Any],
    snapshot: dict[str, Any],
) -> None:
    rows = [
        ("Confidentiality", _format_item_text(item_values.get("system.confidentiality_impact"))),
        ("Integrity", _format_item_text(item_values.get("system.integrity_impact"))),
        ("Availability", _format_item_text(item_values.get("system.availability_impact"))),
    ]
    table = document.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Security Objective"
    header[1].text = "Impact Level"
    for index, (objective, level) in enumerate(rows, start=1):
        row = table.rows[index].cells
        row[0].text = objective
        row[1].text = level

    overall = item_values.get("system.impact_level") or snapshot["profile"]["impact_level"]
    document.add_paragraph(f"Overall FIPS 199 impact level: {_format_item_text(overall)}")
    rationale = item_values.get("system.categorization_rationale")
    if rationale is not None:
        document.add_heading("Categorization Rationale", level=2)
        document.add_paragraph(_format_item_text(rationale))


def _render_control_details_chapter(
    document: DocxDocument,
    snapshot: dict[str, Any],
    item_values: dict[str, Any],
    ordered_controls: list[dict[str, Any]],
) -> None:
    preamble_ids = _ssp_item_ids_for_ref(snapshot, "table1.control-implementation-details")
    control_ids = {item["control_id"] for item in ordered_controls}
    for item_id in preamble_ids:
        if item_id in control_ids:
            continue
        document.add_heading(_item_title(snapshot, item_id), level=2)
        document.add_paragraph(_format_item_text(item_values.get(item_id)))

    for control in ordered_controls:
        document.add_heading(
            f"{control['control_id']} — {control['title']}",
            level=2,
        )
        document.add_paragraph(
            f"Implementation status: {control['implementation_status']}"
        )
        document.add_paragraph(f"Responsibility: {control['responsibility']}")
        document.add_paragraph(
            control["implementation_statement"] or "No statement provided."
        )
        evidence_text = _format_evidence_links(control.get("evidence_links") or [], snapshot)
        if evidence_text:
            document.add_paragraph(f"Evidence: {evidence_text}")


def _render_control_status_table(
    document: DocxDocument,
    ordered_controls: list[dict[str, Any]],
    snapshot: dict[str, Any],
) -> None:
    columns = (
        "Control ID",
        "Title",
        "Implementation Status",
        "Responsibility",
    )
    table = document.add_table(rows=1 + len(ordered_controls), cols=len(columns))
    table.style = "Table Grid"
    for index, label in enumerate(columns):
        table.rows[0].cells[index].text = label
    for row_index, control in enumerate(ordered_controls, start=1):
        cells = table.rows[row_index].cells
        cells[0].text = control["control_id"]
        cells[1].text = control["title"]
        cells[2].text = control["implementation_status"]
        cells[3].text = control["responsibility"]
    if not ordered_controls:
        document.add_paragraph("No controls in profile baseline.")
    document.add_paragraph("")
    document.add_paragraph(
        f"Baseline profile: {snapshot['profile']['profile_id']} "
        f"{snapshot['profile']['version']} ({snapshot['profile']['impact_level']})"
    )


def _ssp_item_ids_for_ref(snapshot: dict[str, Any], requirement_ref: str) -> list[str]:
    ordered: list[str] = []
    for item in snapshot.get("ssp_items") or []:
        refs = item.get("standard_refs") or []
        item_id = item.get("item_id")
        if item_id and requirement_ref in refs and item_id not in ordered:
            ordered.append(item_id)
    return ordered


def _format_evidence_links(links: list[str], snapshot: dict[str, Any]) -> str:
    catalog = snapshot.get("evidence_catalog") or {}
    formatted: list[str] = []
    for link in links:
        artifact_id = link.split(":", 1)[0]
        label = catalog.get(artifact_id)
        if label:
            formatted.append(f"{label} ({link})")
        else:
            formatted.append(link)
    return ", ".join(formatted)


def render_sp800_18_cover(document: DocxDocument, snapshot: dict[str, Any]) -> None:
    """Title page and approval metadata."""
    title = snapshot.get("document_title") or snapshot["system"]["display_name"]
    document.add_heading(f"{title} System Security Plan", level=0)
    document.add_paragraph(SP800_18_SOURCE_LABEL)
    document.add_paragraph(
        "Profile: "
        f"{snapshot['profile']['profile_id']} "
        f"{snapshot['profile']['version']} "
        f"({snapshot['profile']['impact_level']})"
    )
    external_id = snapshot["system"].get("external_system_id")
    if external_id:
        document.add_paragraph(f"System identifier: {external_id}")
    document.add_paragraph(
        f"Approved by {snapshot['approved_by']} at {snapshot['approved_at']}"
    )
    document.add_paragraph(f"Content SHA-256: {snapshot['content_sha256']}")
    document.add_page_break()
