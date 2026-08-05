"""Deterministic JSON and DOCX rendering for approved SSP workspace snapshots."""

from __future__ import annotations

import json
from io import BytesIO
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument

from ato_service.ssp_workspace.sp800_18_docx import (
    nist_control_sort_key,
    render_sp800_18_cover,
    render_sp800_18_docx,
)

EXPORT_SCHEMA_VERSION = "1.1.0"
MAX_SECTIONS = 100
MAX_CONTROLS = 2_000
MAX_QUESTIONS = 2_000
MAX_TEXT_LENGTH = 200_000


class WorkspaceExportValidationError(ValueError):
    """Raised when an approved snapshot cannot be rendered safely."""


def build_workspace_json_export(
    snapshot: dict[str, Any],
    *,
    include_open_questions: bool,
) -> bytes:
    """Return canonical UTF-8 JSON for one approved workspace snapshot."""
    normalized = normalize_export_snapshot(
        snapshot,
        include_open_questions=include_open_questions,
    )
    return (
        json.dumps(
            normalized,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        )
        + "\n"
    ).encode("utf-8")


def build_workspace_docx_export(
    snapshot: dict[str, Any],
    *,
    include_open_questions: bool,
) -> bytes:
    """Return a DOCX rendering of one approved workspace snapshot."""
    normalized = normalize_export_snapshot(
        snapshot,
        include_open_questions=include_open_questions,
    )
    document = Document()
    _set_document_metadata(document, normalized)
    if normalized.get("standard_coverage"):
        render_sp800_18_cover(document, normalized)
        render_sp800_18_docx(document, normalized)
    else:
        _render_header(document, normalized)
        _render_sections(document, normalized["sections"])
        _render_controls(document, normalized["controls"])
    if include_open_questions:
        _render_open_questions(document, normalized["open_questions"])

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def normalize_export_snapshot(
    snapshot: dict[str, Any],
    *,
    include_open_questions: bool,
) -> dict[str, Any]:
    """Validate and normalize the stable export document contract."""
    if not isinstance(snapshot, dict):
        raise WorkspaceExportValidationError("snapshot must be an object")

    workspace_id = _required_text(snapshot, "workspace_id", max_length=128)
    revision_id = _required_text(snapshot, "revision_id", max_length=128)
    content_sha256 = _required_text(snapshot, "content_sha256", max_length=64)
    if len(content_sha256) != 64 or any(
        character not in "0123456789abcdef" for character in content_sha256
    ):
        raise WorkspaceExportValidationError(
            "content_sha256 must be lowercase hexadecimal"
        )
    approved_by = _required_text(snapshot, "approved_by", max_length=255)
    approved_at = _required_text(snapshot, "approved_at", max_length=64)

    system = _required_object(snapshot, "system")
    profile = _required_object(snapshot, "profile")
    sections = _normalize_sections(snapshot.get("sections"))
    controls = _normalize_controls(
        snapshot.get("controls"),
        control_order=snapshot.get("control_order"),
    )
    questions = _normalize_questions(snapshot.get("questions"))
    open_questions = (
        [question for question in questions if question["status"] == "open"]
        if include_open_questions
        else []
    )

    document_title = _optional_text(snapshot, "document_title", max_length=500)
    facts = _normalize_facts(snapshot.get("facts"))
    standard_coverage = _normalize_standard_coverage(snapshot.get("standard_coverage"))
    ssp_items = _normalize_ssp_items(snapshot.get("ssp_items"))
    control_order = _normalize_control_order(snapshot.get("control_order"))
    evidence_catalog = _normalize_evidence_catalog(snapshot.get("evidence_catalog"))

    normalized: dict[str, Any] = {
        "schema_version": EXPORT_SCHEMA_VERSION,
        "workspace_id": workspace_id,
        "revision_id": revision_id,
        "content_sha256": content_sha256,
        "approved_by": approved_by,
        "approved_at": approved_at,
        "system": {
            "display_name": _required_text(system, "display_name", max_length=255),
            "external_system_id": _optional_text(
                system,
                "external_system_id",
                max_length=255,
            ),
        },
        "profile": {
            "profile_id": _required_text(profile, "profile_id", max_length=128),
            "version": _required_text(profile, "version", max_length=64),
            "impact_level": _required_text(
                profile,
                "impact_level",
                max_length=16,
            ),
        },
        "sections": sections,
        "controls": controls,
        "open_questions": open_questions,
    }
    if document_title is not None:
        normalized["document_title"] = document_title
    if facts:
        normalized["facts"] = facts
    if standard_coverage:
        normalized["standard_coverage"] = standard_coverage
    if ssp_items:
        normalized["ssp_items"] = ssp_items
    if control_order:
        normalized["control_order"] = control_order
    if evidence_catalog:
        normalized["evidence_catalog"] = evidence_catalog
    return normalized


def _normalize_sections(value: Any) -> list[dict[str, Any]]:
    items = _bounded_object_list(value, field_name="sections", maximum=MAX_SECTIONS)
    normalized = [
        {
            "section_id": _required_text(item, "section_id", max_length=128),
            "title": _required_text(item, "title", max_length=500),
            "order": _required_integer(item, "order", minimum=0),
            "state": _required_text(item, "state", max_length=32),
            "content": _required_text(
                item,
                "content",
                max_length=MAX_TEXT_LENGTH,
                allow_empty=True,
            ),
        }
        for item in items
    ]
    section_ids = [item["section_id"] for item in normalized]
    _require_unique(section_ids, field_name="section_id")
    return sorted(normalized, key=lambda item: (item["order"], item["section_id"]))


def _normalize_controls(
    value: Any,
    *,
    control_order: Any = None,
) -> list[dict[str, Any]]:
    items = _bounded_object_list(value, field_name="controls", maximum=MAX_CONTROLS)
    normalized: list[dict[str, Any]] = []
    for item in items:
        evidence_links = item.get("evidence_links", [])
        if not isinstance(evidence_links, list) or any(
            not isinstance(link, str) or not link.strip()
            for link in evidence_links
        ):
            raise WorkspaceExportValidationError(
                "controls evidence_links must contain non-empty strings"
            )
        normalized.append(
            {
                "control_id": _required_text(item, "control_id", max_length=128),
                "title": _required_text(item, "title", max_length=500),
                "state": _required_text(item, "state", max_length=32),
                "implementation_status": _required_text(
                    item,
                    "implementation_status",
                    max_length=64,
                ),
                "responsibility": _required_text(
                    item,
                    "responsibility",
                    max_length=64,
                ),
                "implementation_statement": _required_text(
                    item,
                    "implementation_statement",
                    max_length=MAX_TEXT_LENGTH,
                    allow_empty=True,
                ),
                "evidence_links": sorted(set(evidence_links)),
            }
        )
    control_ids = [item["control_id"] for item in normalized]
    _require_unique(control_ids, field_name="control_id")
    order = _normalize_control_order(control_order)
    if order:
        by_id = {item["control_id"]: item for item in normalized}
        ordered = [by_id[cid] for cid in order if cid in by_id]
        seen = {item["control_id"] for item in ordered}
        ordered.extend(
            sorted(
                (item for item in normalized if item["control_id"] not in seen),
                key=lambda item: nist_control_sort_key(item["control_id"]),
            )
        )
        return ordered
    return sorted(normalized, key=lambda item: nist_control_sort_key(item["control_id"]))


def _normalize_questions(value: Any) -> list[dict[str, Any]]:
    items = _bounded_object_list(value, field_name="questions", maximum=MAX_QUESTIONS)
    normalized = [
        {
            "question_id": _required_text(item, "question_id", max_length=128),
            "target": _required_text(item, "target", max_length=500),
            "question": _required_text(item, "question", max_length=8_000),
            "owner_type": _required_text(item, "owner_type", max_length=64),
            "status": _required_text(item, "status", max_length=32),
        }
        for item in items
    ]
    question_ids = [item["question_id"] for item in normalized]
    _require_unique(question_ids, field_name="question_id")
    return sorted(normalized, key=lambda item: item["question_id"])


def _set_document_metadata(
    document: DocxDocument,
    snapshot: dict[str, Any],
) -> None:
    title_name = snapshot.get("document_title") or snapshot["system"]["display_name"]
    document.core_properties.title = f"{title_name} System Security Plan"
    document.core_properties.subject = "System Security Plan"
    document.core_properties.author = snapshot["approved_by"]
    document.core_properties.keywords = (
        f"{snapshot['profile']['profile_id']},"
        f"{snapshot['profile']['version']},"
        f"{snapshot['content_sha256']}"
    )


def _render_header(
    document: DocxDocument,
    snapshot: dict[str, Any],
) -> None:
    document.add_heading(
        f"{snapshot['system']['display_name']} System Security Plan",
        level=0,
    )
    document.add_paragraph(
        "Profile: "
        f"{snapshot['profile']['profile_id']} "
        f"{snapshot['profile']['version']} "
        f"({snapshot['profile']['impact_level']})"
    )
    document.add_paragraph(
        f"Approved by {snapshot['approved_by']} at {snapshot['approved_at']}"
    )
    document.add_paragraph(f"Content SHA-256: {snapshot['content_sha256']}")


def _render_sections(
    document: DocxDocument,
    sections: list[dict[str, Any]],
) -> None:
    for section in sections:
        document.add_heading(section["title"], level=1)
        document.add_paragraph(section["content"] or "No content provided.")


def _render_controls(
    document: DocxDocument,
    controls: list[dict[str, Any]],
) -> None:
    document.add_heading("Security Controls", level=1)
    for control in controls:
        document.add_heading(
            f"{control['control_id']} — {control['title']}",
            level=2,
        )
        document.add_paragraph(
            f"Implementation status: {control['implementation_status']}"
        )
        document.add_paragraph(f"Responsibility: {control['responsibility']}")
        document.add_paragraph(
            control["implementation_statement"] or "No statement provided."
        )
        if control["evidence_links"]:
            document.add_paragraph(
                "Evidence: " + ", ".join(control["evidence_links"])
            )


def _render_open_questions(
    document: DocxDocument,
    questions: list[dict[str, Any]],
) -> None:
    if not questions:
        return
    document.add_heading("Unresolved Items", level=1)
    for question in questions:
        document.add_paragraph(
            f"{question['target']}: {question['question']} "
            f"(owner: {question['owner_type']})",
            style="List Bullet",
        )


def _bounded_object_list(
    value: Any,
    *,
    field_name: str,
    maximum: int,
) -> list[dict[str, Any]]:
    if not isinstance(value, list):
        raise WorkspaceExportValidationError(f"{field_name} must be an array")
    if len(value) > maximum:
        raise WorkspaceExportValidationError(
            f"{field_name} exceeds the maximum of {maximum}"
        )
    if any(not isinstance(item, dict) for item in value):
        raise WorkspaceExportValidationError(
            f"{field_name} must contain objects"
        )
    return value


def _required_object(document: dict[str, Any], field_name: str) -> dict[str, Any]:
    value = document.get(field_name)
    if not isinstance(value, dict):
        raise WorkspaceExportValidationError(f"{field_name} must be an object")
    return value


def _required_text(
    document: dict[str, Any],
    field_name: str,
    *,
    max_length: int,
    allow_empty: bool = False,
) -> str:
    value = document.get(field_name)
    if not isinstance(value, str):
        raise WorkspaceExportValidationError(f"{field_name} must be a string")
    normalized = value.strip()
    if not normalized and not allow_empty:
        raise WorkspaceExportValidationError(f"{field_name} must not be empty")
    if len(value) > max_length:
        raise WorkspaceExportValidationError(
            f"{field_name} exceeds the maximum length of {max_length}"
        )
    return value


def _optional_text(
    document: dict[str, Any],
    field_name: str,
    *,
    max_length: int,
) -> str | None:
    value = document.get(field_name)
    if value is None:
        return None
    if not isinstance(value, str) or len(value) > max_length:
        raise WorkspaceExportValidationError(
            f"{field_name} must be null or a string of at most {max_length} characters"
        )
    return value


def _required_integer(
    document: dict[str, Any],
    field_name: str,
    *,
    minimum: int,
) -> int:
    value = document.get(field_name)
    if isinstance(value, bool) or not isinstance(value, int) or value < minimum:
        raise WorkspaceExportValidationError(
            f"{field_name} must be an integer greater than or equal to {minimum}"
        )
    return value


def _require_unique(values: list[str], *, field_name: str) -> None:
    if len(values) != len(set(values)):
        raise WorkspaceExportValidationError(
            f"duplicate {field_name} values are not allowed"
        )


def _normalize_facts(value: Any) -> dict[str, Any]:
    if value is None:
        return {}
    if not isinstance(value, dict):
        raise WorkspaceExportValidationError("facts must be an object")
    if len(value) > 500:
        raise WorkspaceExportValidationError("facts exceeds the maximum of 500")
    normalized: dict[str, Any] = {}
    for key, fact_value in value.items():
        if not isinstance(key, str) or not key.strip() or len(key) > 255:
            raise WorkspaceExportValidationError("facts keys must be non-empty strings")
        normalized[key] = fact_value
    return dict(sorted(normalized.items(), key=lambda item: item[0]))


def _normalize_standard_coverage(value: Any) -> list[dict[str, Any]]:
    if value is None:
        return []
    if not isinstance(value, list):
        raise WorkspaceExportValidationError("standard_coverage must be an array")
    if len(value) > 100:
        raise WorkspaceExportValidationError(
            "standard_coverage exceeds the maximum of 100"
        )
    normalized: list[dict[str, Any]] = []
    for index, item in enumerate(value):
        if not isinstance(item, dict):
            raise WorkspaceExportValidationError(
                "standard_coverage must contain objects"
            )
        normalized.append(
            {
                "source_id": _required_text(item, "source_id", max_length=128),
                "requirement_id": _required_text(
                    item,
                    "requirement_id",
                    max_length=256,
                ),
                "title": _required_text(item, "title", max_length=500),
                "coverage_kind": _required_text(
                    item,
                    "coverage_kind",
                    max_length=32,
                ),
                "item_ids": _normalize_string_list(
                    item.get("item_ids"),
                    field_name=f"standard_coverage[{index}].item_ids",
                    maximum=100,
                ),
                "required": bool(item.get("required", True)),
            }
        )
    return normalized


def _normalize_ssp_items(value: Any) -> list[dict[str, Any]]:
    if value is None:
        return []
    if not isinstance(value, list):
        raise WorkspaceExportValidationError("ssp_items must be an array")
    if len(value) > 500:
        raise WorkspaceExportValidationError("ssp_items exceeds the maximum of 500")
    normalized: list[dict[str, Any]] = []
    for item in value:
        if not isinstance(item, dict):
            raise WorkspaceExportValidationError("ssp_items must contain objects")
        normalized.append(
            {
                "item_id": _required_text(item, "item_id", max_length=255),
                "title": _required_text(item, "title", max_length=500),
                "value_type": _required_text(item, "value_type", max_length=32),
                "standard_refs": _normalize_string_list(
                    item.get("standard_refs"),
                    field_name="ssp_items.standard_refs",
                    maximum=32,
                ),
            }
        )
    return normalized


def _normalize_control_order(value: Any) -> list[str]:
    return _normalize_string_list(
        value,
        field_name="control_order",
        maximum=MAX_CONTROLS,
    )


def _normalize_evidence_catalog(value: Any) -> dict[str, str]:
    if value is None:
        return {}
    if not isinstance(value, dict):
        raise WorkspaceExportValidationError("evidence_catalog must be an object")
    if len(value) > 5_000:
        raise WorkspaceExportValidationError(
            "evidence_catalog exceeds the maximum of 5000"
        )
    normalized: dict[str, str] = {}
    for key, label in value.items():
        if not isinstance(key, str) or len(key) > 128:
            raise WorkspaceExportValidationError(
                "evidence_catalog keys must be strings of at most 128 characters"
            )
        if not isinstance(label, str) or not label.strip() or len(label) > 255:
            raise WorkspaceExportValidationError(
                "evidence_catalog values must be non-empty strings"
            )
        normalized[key] = label.strip()
    return dict(sorted(normalized.items(), key=lambda item: item[0]))


def _normalize_string_list(
    value: Any,
    *,
    field_name: str,
    maximum: int,
) -> list[str]:
    if value is None:
        return []
    if not isinstance(value, list):
        raise WorkspaceExportValidationError(f"{field_name} must be an array")
    if len(value) > maximum:
        raise WorkspaceExportValidationError(
            f"{field_name} exceeds the maximum of {maximum}"
        )
    normalized: list[str] = []
    for item in value:
        if not isinstance(item, str) or not item.strip():
            raise WorkspaceExportValidationError(
                f"{field_name} must contain non-empty strings"
            )
        if len(item) > 255:
            raise WorkspaceExportValidationError(
                f"{field_name} entries must be at most 255 characters"
            )
        normalized.append(item.strip())
    return normalized
