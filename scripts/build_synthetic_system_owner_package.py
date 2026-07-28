#!/usr/bin/env python3
"""Build compact pre-ATO synthetic system-owner evidence for SSP evaluation."""

from __future__ import annotations

import hashlib
import json
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = (
    ROOT
    / "data"
    / "synthetic-packages"
    / "fisma-rev5-pre-ato-system-owner"
)
ZIP_PATH = OUTPUT_DIR / "fgrs-system-owner-artifacts.zip"

NAVY = "#17365D"
BLUE = "#D9EAF7"
GREEN = "#DDEEDC"
ORANGE = "#FCE4C5"
GRAY = "#F1F3F5"
RED = "#F8D7DA"
WHITE = "#FFFFFF"
TEXT = "#17212B"


def _font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = (
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")
        if bold
        else Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf")
        if bold
        else Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf"),
    )
    for candidate in candidates:
        if candidate.is_file():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def _wrapped(draw: ImageDraw.ImageDraw, text: str, width: int, font) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=font)[2] <= width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def _box(
    draw: ImageDraw.ImageDraw,
    bounds: tuple[int, int, int, int],
    label: str,
    *,
    fill: str,
    outline: str = NAVY,
    font_size: int = 25,
) -> None:
    x1, y1, x2, y2 = bounds
    draw.rounded_rectangle(bounds, radius=12, fill=fill, outline=outline, width=3)
    font = _font(font_size, bold=True)
    lines = _wrapped(draw, label, x2 - x1 - 24, font)
    line_height = font_size + 6
    y = y1 + ((y2 - y1) - line_height * len(lines)) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text(
            (x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2, y),
            line,
            fill=TEXT,
            font=font,
        )
        y += line_height


def _arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    *,
    label: str | None = None,
) -> None:
    draw.line((start, end), fill=NAVY, width=4)
    x2, y2 = end
    x1, y1 = start
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - 14 * direction, y2 - 8), (x2 - 14 * direction, y2 + 8)]
    else:
        direction = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 8, y2 - 14 * direction), (x2 + 8, y2 - 14 * direction)]
    draw.polygon(points, fill=NAVY)
    if label:
        font = _font(18)
        bbox = draw.textbbox((0, 0), label, font=font)
        cx = (x1 + x2) / 2
        cy = (y1 + y2) / 2
        pad = 5
        draw.rectangle(
            (
                cx - (bbox[2] - bbox[0]) / 2 - pad,
                cy - 14,
                cx + (bbox[2] - bbox[0]) / 2 + pad,
                cy + 14,
            ),
            fill=WHITE,
        )
        draw.text(
            (cx - (bbox[2] - bbox[0]) / 2, cy - 11),
            label,
            fill=TEXT,
            font=font,
        )


def _canvas(title: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (1800, 1100), WHITE)
    draw = ImageDraw.Draw(image)
    title_font = _font(40, bold=True)
    draw.text((70, 42), title, fill=NAVY, font=title_font)
    draw.line((70, 100, 1730, 100), fill=NAVY, width=4)
    return image, draw


def _architecture_diagram(path: Path) -> None:
    image, draw = _canvas("FGRS Production Architecture")
    boundary = (330, 150, 1480, 965)
    draw.rounded_rectangle(boundary, radius=20, fill="#F8FBFD", outline=NAVY, width=5)
    draw.text(
        (365, 168),
        "Agency Enterprise Cloud - FGRS Production Environment",
        fill=NAVY,
        font=_font(26, bold=True),
    )

    boxes = {
        "users": (60, 320, 270, 440),
        "idp": (60, 570, 270, 690),
        "web": (410, 300, 660, 430),
        "api": (760, 300, 1010, 430),
        "worker": (760, 585, 1010, 715),
        "db": (1130, 260, 1370, 390),
        "storage": (1130, 490, 1370, 620),
        "keys": (1130, 720, 1370, 850),
        "siem": (1530, 260, 1740, 380),
        "email": (1530, 485, 1740, 605),
        "backup": (1530, 710, 1740, 830),
    }
    labels = {
        "users": ("Agency Staff", ORANGE),
        "idp": ("Agency Identity Provider", GREEN),
        "web": ("Internal Web Application", BLUE),
        "api": ("Application API", BLUE),
        "worker": ("Background Worker", BLUE),
        "db": ("Managed PostgreSQL", GRAY),
        "storage": ("Encrypted Object Storage", GRAY),
        "keys": ("Agency Key Service", GREEN),
        "siem": ("Agency SIEM", GREEN),
        "email": ("Agency Email Relay", GREEN),
        "backup": ("Backup Vault", GREEN),
    }
    for key, bounds in boxes.items():
        _box(draw, bounds, labels[key][0], fill=labels[key][1], font_size=22)

    _arrow(draw, (270, 380), (410, 365), label="HTTPS")
    _arrow(draw, (270, 630), (410, 405), label="OIDC")
    _arrow(draw, (660, 365), (760, 365), label="HTTPS")
    _arrow(draw, (1010, 345), (1130, 325), label="TLS")
    _arrow(draw, (1010, 390), (1130, 555), label="HTTPS")
    _arrow(draw, (885, 430), (885, 585), label="Queue")
    _arrow(draw, (1010, 650), (1130, 555))
    _arrow(draw, (1370, 325), (1530, 320), label="Logs")
    _arrow(draw, (1370, 555), (1530, 545), label="Mail")
    _arrow(draw, (1370, 785), (1530, 770), label="Backup")
    draw.text(
        (70, 1020),
        "Synthetic engineering diagram. Components outside the cloud box are managed shared services.",
        fill=TEXT,
        font=_font(20),
    )
    image.save(path, format="PNG", optimize=True)


def _data_flow_diagram(path: Path) -> None:
    image, draw = _canvas("FGRS Business Data Flow")
    stages = [
        ((70, 250, 310, 390), "Program Reviewer", ORANGE),
        ((390, 250, 630, 390), "Internal Web Application", BLUE),
        ((710, 250, 950, 390), "Application API", BLUE),
        ((1030, 170, 1270, 310), "PostgreSQL Records", GRAY),
        ((1030, 410, 1270, 550), "Supporting Documents", GRAY),
        ((1440, 170, 1690, 310), "Agency SIEM", GREEN),
        ((1440, 410, 1690, 550), "Agency Email Relay", GREEN),
    ]
    for bounds, label, fill in stages:
        _box(draw, bounds, label, fill=fill, font_size=22)
    _arrow(draw, (310, 320), (390, 320), label="TLS 1.2+")
    _arrow(draw, (630, 320), (710, 320), label="JSON API")
    _arrow(draw, (950, 285), (1030, 240), label="SQL/TLS")
    _arrow(draw, (950, 350), (1030, 480), label="HTTPS")
    _arrow(draw, (1270, 240), (1440, 240), label="Audit events")
    _arrow(draw, (1270, 480), (1440, 480), label="Notifications")

    _box(draw, (220, 690, 520, 830), "Nightly Database Backup", fill=GREEN, font_size=22)
    _box(draw, (750, 690, 1050, 830), "Encrypted Backup Transfer", fill=BLUE, font_size=22)
    _box(draw, (1280, 690, 1580, 830), "Agency Backup Vault", fill=GREEN, font_size=22)
    _arrow(draw, (520, 760), (750, 760), label="Encrypted")
    _arrow(draw, (1050, 760), (1280, 760), label="Private endpoint")

    draw.text(
        (70, 940),
        "Data handled: grant application records, business contact information, reviewer notes, and attachments.",
        fill=TEXT,
        font=_font(21),
    )
    draw.text(
        (70, 980),
        "No direct public access. Administrative access uses the agency privileged access service.",
        fill=TEXT,
        font=_font(21),
    )
    image.save(path, format="PNG", optimize=True)


def _operations_diagram(path: Path) -> None:
    image, draw = _canvas("FGRS Operational Monitoring and Response")
    lanes = (
        ("Identity and access", 170, BLUE),
        ("Logging and monitoring", 450, GREEN),
        ("Incident handling", 730, RED),
    )
    labels = (
        (
            "Manager Request",
            "Data Owner Approval",
            "Role Assignment",
            "Quarterly Review",
            "Disable Access",
        ),
        (
            "Application Events",
            "Log Forwarder",
            "Agency SIEM",
            "24x7 Alerting",
            "Daily Review",
        ),
        (
            "Alert or Report",
            "Security Operations",
            "Notify Owner",
            "Contain and Recover",
            "Post-Incident Review",
        ),
    )
    for lane_index, (lane, y, fill) in enumerate(lanes):
        draw.text((70, y - 42), lane, fill=NAVY, font=_font(24, bold=True))
        x_positions = (70, 410, 750, 1090, 1430)
        for index, x in enumerate(x_positions):
            _box(
                draw,
                (x, y, x + 270, y + 125),
                labels[lane_index][index],
                fill=fill,
                font_size=21,
            )
            if index < len(x_positions) - 1:
                _arrow(draw, (x + 270, y + 62), (x_positions[index + 1], y + 62))
    draw.text(
        (70, 1010),
        "Synthetic operations diagram derived from the FGRS team runbook.",
        fill=TEXT,
        font=_font(20),
    )
    image.save(path, format="PNG", optimize=True)


def _set_cell_shading(cell, color: str) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color.lstrip("#"))
    tc_pr.append(shading)


def _document(title: str, subtitle: str) -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    for name, size, color in (
        ("Title", 24, NAVY),
        ("Heading 1", 16, NAVY),
        ("Heading 2", 12, NAVY),
    ):
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    heading = document.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = document.add_paragraph(subtitle)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    document.add_paragraph("Document status: Synthetic operational artifact for product testing")
    return document


def _add_table(document: Document, headers: tuple[str, ...], rows: list[tuple[str, ...]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        _set_cell_shading(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def _add_bullets(document: Document, values: tuple[str, ...]) -> None:
    for value in values:
        document.add_paragraph(value, style="List Bullet")


def _build_product_overview(path: Path, architecture_path: Path, data_flow_path: Path) -> None:
    doc = _document(
        "Federal Grants Review System",
        "Product overview, data inventory, and production architecture",
    )
    doc.add_heading("1. Product summary", level=1)
    doc.add_paragraph(
        "The Federal Grants Review System (FGRS) is an internal agency web "
        "application used by program-office staff to receive, review, score, "
        "and track grant application packages after submission through a "
        "separate public intake service. FGRS does not accept direct public "
        "connections. The production system identifier is FGRS-PRD."
    )
    _add_table(
        doc,
        ("Attribute", "Operational value"),
        [
            ("Business owner", "Dana Holloway, Director, Office of Grants Operations"),
            ("Technical owner", "Riley Chen, Application Platform Manager"),
            ("Support hours", "06:00-22:00 Eastern, Monday through Friday"),
            ("Availability objective", "99.9 percent monthly during support hours"),
            ("Hosting", "Agency enterprise cloud, agency-managed production account"),
            ("Environments", "Development, test, and production are separate"),
            ("Production region", "Primary agency east region; backup vault in secondary region"),
        ],
    )

    doc.add_heading("2. Users and business functions", level=1)
    _add_bullets(
        doc,
        (
            "Program Reviewers view assigned applications, record scores, and enter review notes.",
            "Program Managers assign work, approve final recommendations, and run reports.",
            "Records Staff export approved records for transfer to the agency records repository.",
            "Read-only Auditors can view records and audit history but cannot change case data.",
            "Application Administrators manage application roles and reference data.",
            "Platform Administrators maintain the runtime platform through the agency privileged access service.",
        ),
    )

    doc.add_heading("3. Information handled", level=1)
    _add_table(
        doc,
        ("Information", "Examples", "Handling"),
        [
            (
                "Grant application records",
                "Narratives, budgets, eligibility documents, scoring results",
                "Agency internal; some records may be pre-decisional",
            ),
            (
                "Business contact information",
                "Names, work email, work phone, organization",
                "Limited business contact information",
            ),
            (
                "Reviewer information",
                "Agency username, assignments, comments, approvals",
                "Agency internal",
            ),
            (
                "Supporting attachments",
                "PDF, DOCX, and XLSX files transferred from the intake service",
                "Stored encrypted; malware-scanned before transfer",
            ),
            (
                "Operational records",
                "Audit events, application logs, job status, performance metrics",
                "Restricted to operations and security personnel",
            ),
        ],
    )
    doc.add_paragraph(
        "The product team treats production confidentiality, integrity, and "
        "availability as Moderate. The system is not approved to store "
        "classified information, payment card data, protected health "
        "information, or applicant bank-account credentials."
    )

    doc.add_heading("4. Production boundary and components", level=1)
    doc.add_picture(str(architecture_path), width=Inches(7.1))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "The FGRS production environment includes the internal web application, "
        "application API, background worker, managed PostgreSQL database, "
        "encrypted object storage, deployment namespace, workload identities, "
        "network policies, and application-specific encryption keys. The agency "
        "identity provider, SIEM, email relay, vulnerability scanner, privileged "
        "access service, and backup vault are shared services operated outside "
        "the product environment."
    )

    doc.add_heading("5. Data flow and external services", level=1)
    doc.add_picture(str(data_flow_path), width=Inches(7.1))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_table(
        doc,
        ("Service", "Direction and purpose", "Protection"),
        [
            (
                "Agency identity provider",
                "Inbound authentication and group claims",
                "OIDC over TLS; phishing-resistant MFA using agency PIV",
            ),
            (
                "Public intake service",
                "Inbound transfer of accepted application packages",
                "Mutual TLS, workload identity, schema validation, malware-scan status required",
            ),
            (
                "Agency SIEM",
                "Outbound audit and security events",
                "TLS log forwarding through the central collector",
            ),
            (
                "Agency email relay",
                "Outbound workflow notifications",
                "Authenticated TLS connection; no attachments are sent",
            ),
            (
                "Vulnerability service",
                "Inbound authenticated scanning of approved endpoints",
                "Dedicated scanner identity and private network path",
            ),
            (
                "Backup vault",
                "Outbound encrypted database and object backups",
                "Private endpoint and separate agency-managed encryption key",
            ),
        ],
    )
    doc.save(path)


def _build_operations_guide(path: Path, operations_path: Path) -> None:
    doc = _document(
        "FGRS Production Operations Guide",
        "Version 3.4 - maintained by the application platform team",
    )
    doc.add_heading("1. Deployment and configuration", level=1)
    doc.add_paragraph(
        "Production changes are deployed from the approved source repository "
        "through the agency continuous-integration service. Builds create signed "
        "container images, scan dependencies and image layers, and publish "
        "approved images to the agency registry. Production deployment requires "
        "a peer-reviewed change request and approval from the technical owner. "
        "Direct interactive changes to running containers are prohibited."
    )
    _add_bullets(
        doc,
        (
            "A standard monthly maintenance window is used for routine platform and application updates.",
            "Critical remotely exploitable defects are evaluated within one business day and remediated or mitigated within 15 calendar days.",
            "High-severity defects are remediated within 30 calendar days.",
            "Infrastructure and application configuration are version-controlled; protected branches require two reviewers.",
            "Secrets are stored in the agency secret-management service and are not stored in source code, container images, or deployment manifests.",
            "Production and non-production use separate accounts, networks, databases, keys, and workload identities.",
        ),
    )

    doc.add_heading("2. Logging and monitoring", level=1)
    doc.add_picture(str(operations_path), width=Inches(7.1))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_table(
        doc,
        ("Event source", "Recorded details", "Operational use"),
        [
            (
                "Authentication",
                "Success, failure, user, source, time, session identifier",
                "Brute-force and unusual-access alerting",
            ),
            (
                "Authorization",
                "Role changes, denied actions, privileged actions",
                "Access review and misuse detection",
            ),
            (
                "Business records",
                "Create, update, assignment, score, approval, export",
                "Accountability and troubleshooting",
            ),
            (
                "Platform",
                "Deployment, container restart, health, database connection, job failure",
                "Availability and reliability monitoring",
            ),
            (
                "Security tools",
                "Image findings, endpoint findings, dependency findings",
                "Remediation tracking",
            ),
        ],
    )
    doc.add_paragraph(
        "Application audit events are forwarded to the agency SIEM within five "
        "minutes. Security Operations provides 24x7 alerting. The application "
        "team reviews service health and failed jobs each business day. Logs "
        "remain searchable for 90 days and are archived for one year. "
        "Application logs exclude access tokens, passwords, session cookies, "
        "full document contents, and sensitive form values."
    )

    doc.add_heading("3. Backup and recovery", level=1)
    _add_bullets(
        doc,
        (
            "Managed database backups run nightly and transaction logs support point-in-time recovery.",
            "Object-storage versioning is enabled and deleted objects are retained for 35 days.",
            "Nightly backups are retained for 35 days; monthly recovery points are retained for one year.",
            "Backups use a separate agency-managed encryption key and are copied to the secondary-region backup vault.",
            "The team performs a documented restore test each quarter.",
            "The recovery objective is eight hours and the maximum acceptable data loss is 24 hours.",
            "Recovery requires the incident commander or continuity coordinator to authorize failover.",
        ),
    )

    doc.add_heading("4. Vulnerability and maintenance activities", level=1)
    _add_bullets(
        doc,
        (
            "Container images and application dependencies are scanned on every build and at least weekly.",
            "The agency scanning service performs an authenticated infrastructure scan monthly.",
            "The product team reviews new findings each business day and records remediation owners and due dates in the engineering ticket system.",
            "Annual application penetration testing is coordinated by the agency security testing team.",
            "Unsupported software versions are blocked from production deployment by pipeline policy.",
            "Time is synchronized from the agency time service.",
        ),
    )

    doc.add_heading("5. Incident and service-event handling", level=1)
    doc.add_paragraph(
        "Users report suspected security events through the agency service desk "
        "or security hotline. Automated alerts are routed to Security Operations. "
        "Security Operations opens the incident record, performs initial triage, "
        "and contacts the FGRS technical owner and cybersecurity liaison. The "
        "application team preserves relevant logs, isolates affected workloads "
        "when directed, rotates exposed credentials, restores trusted images, "
        "and validates service health. External reporting and notification are "
        "handled under the agency incident-response process. A post-incident "
        "review records root cause, corrective actions, and assigned owners."
    )
    doc.save(path)


def _build_access_procedure(path: Path) -> None:
    doc = _document(
        "FGRS User Access and Support Procedures",
        "Daily operating procedure for program and platform personnel",
    )
    doc.add_heading("1. Standard user access", level=1)
    _add_bullets(
        doc,
        (
            "The user's manager submits an access request in the agency service desk.",
            "The request identifies the business role, program office, and required grant portfolio.",
            "The program data owner approves Reviewer, Program Manager, Records Staff, and Auditor access.",
            "An Application Administrator assigns the approved role after confirming the user's active agency identity.",
            "Users authenticate through the agency identity provider with PIV-based multifactor authentication.",
            "The application denies access when required group claims are absent or the account is disabled.",
        ),
    )

    doc.add_heading("2. Privileged access", level=1)
    doc.add_paragraph(
        "Application Administrator and Platform Administrator access requires a "
        "separate privileged account, technical-owner approval, and enrollment "
        "in the agency privileged access service. Privileged sessions are "
        "initiated through the managed access path and are logged. Platform "
        "Administrators do not receive standing database credentials. Emergency "
        "access is time-limited, requires an incident or change ticket, and is "
        "reviewed the next business day."
    )

    doc.add_heading("3. Reviews, transfers, and removal", level=1)
    _add_bullets(
        doc,
        (
            "Application Administrators export active users and roles each quarter.",
            "Program Managers confirm business access; the technical owner confirms privileged access.",
            "Unconfirmed or unnecessary access is removed within one business day after review.",
            "The identity provider disables separated personnel through the agency offboarding feed.",
            "The application team removes application roles within four hours of a high-priority separation notice.",
            "Transferred personnel require a new manager and data-owner approval before retaining access.",
            "Inactive application role assignments are removed after 60 days unless the manager documents a continuing need.",
        ),
    )

    doc.add_heading("4. Session and authentication behavior", level=1)
    _add_bullets(
        doc,
        (
            "Sessions end after 15 minutes of inactivity and after 12 hours regardless of activity.",
            "Five failed authentication attempts within 15 minutes trigger a 30-minute application lockout while the agency identity provider applies its own protections.",
            "Concurrent sessions are limited to two per user.",
            "The application displays the agency-authorized-use notice before access.",
            "Service-to-service identities use short-lived workload credentials and cannot perform interactive login.",
        ),
    )

    doc.add_heading("5. Support and escalation", level=1)
    _add_table(
        doc,
        ("Issue", "Initial destination", "Escalation"),
        [
            ("Access request or role issue", "Agency service desk", "Application Administrator"),
            ("Application availability", "Agency service desk", "Platform on-call engineer"),
            ("Suspected security event", "Security Operations", "Technical owner and cybersecurity liaison"),
            ("Data-quality or workflow issue", "Program support queue", "Program Manager"),
            ("Restore request", "Platform support queue", "Technical owner approval required"),
        ],
    )
    doc.add_paragraph(
        "Named contacts and phone numbers are maintained in the agency service "
        "directory rather than in this procedure."
    )
    doc.save(path)


def _write_configuration_export(path: Path) -> None:
    document = {
        "application": {
            "name": "Federal Grants Review System",
            "short_name": "FGRS",
            "environment": "production",
            "system_identifier": "FGRS-PRD",
            "release": "3.4.2",
        },
        "hosting": {
            "platform": "agency_enterprise_cloud",
            "deployment_type": "managed_kubernetes",
            "network_exposure": "agency_internal_only",
            "replicas": {"web": 3, "api": 3, "worker": 2},
        },
        "session": {
            "idle_timeout_minutes": 15,
            "absolute_timeout_hours": 12,
            "maximum_concurrent_sessions": 2,
            "failed_attempt_threshold": 5,
            "lockout_minutes": 30,
        },
        "encryption": {
            "minimum_tls": "1.2",
            "database_at_rest": "agency_managed_key",
            "object_storage_at_rest": "agency_managed_key",
            "backup_at_rest": "separate_agency_managed_key",
        },
        "logging": {
            "destination": "agency_siem",
            "forwarding_target_minutes": 5,
            "searchable_days": 90,
            "archive_days": 365,
            "sensitive_value_logging": False,
        },
        "backup": {
            "database_frequency": "nightly",
            "point_in_time_recovery": True,
            "daily_retention_days": 35,
            "monthly_retention_months": 12,
            "restore_test_frequency": "quarterly",
            "recovery_time_hours": 8,
            "recovery_point_hours": 24,
        },
        "integrations": [
            "agency_identity_provider",
            "application_intake_service",
            "agency_siem",
            "agency_email_relay",
            "agency_vulnerability_scanner",
            "agency_backup_vault",
        ],
        "secrets_included": False,
        "synthetic": True,
    }
    path.write_text(json.dumps(document, indent=2, sort_keys=True) + "\n", encoding="utf-8")


def _write_inventory(path: Path) -> None:
    rows = [
        ("fgrs-web", "Web application", "3.4.2", "3", "Internal ingress"),
        ("fgrs-api", "Application API", "3.4.2", "3", "Private service"),
        ("fgrs-worker", "Background processing", "3.4.2", "2", "Private service"),
        ("postgresql", "Managed relational database", "16", "HA pair", "Private subnet"),
        ("object-storage", "Supporting document storage", "Managed", "Regional", "Private endpoint"),
        ("central-log-forwarder", "Audit and operational log forwarding", "Agency managed", "2", "Shared service"),
        ("secret-service", "Runtime secret delivery", "Agency managed", "Regional", "Shared service"),
    ]
    lines = ["Component | Purpose | Version | Quantity | Connectivity"]
    lines.extend(" | ".join(row) for row in rows)
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def _write_readme(path: Path) -> None:
    path.write_text(
        """# FGRS synthetic system-owner artifact package

This package represents material a system owner and engineering team could
reasonably provide before beginning authorization documentation.

It intentionally contains:

- no SSP;
- no control identifiers or control mappings;
- no implementation statements written for authorization;
- no authorization decision or compliance claim.

Recommended evaluation upload set:

1. `01-fgrs-product-and-architecture.docx`
2. `02-fgrs-production-operations-guide.docx`
3. `03-fgrs-user-access-procedures.docx`
4. `04-fgrs-production-config.json`
5. `05-fgrs-component-inventory.txt`
6. the three PNG diagrams

The documents and diagrams overlap in realistic ways. The agent should only
populate SSP content and control statements supported by the artifacts. It
should leave unsupported details unknown and create targeted questions.

All people, systems, addresses, and operational details are synthetic.
""",
        encoding="utf-8",
    )


def _write_manifest(path: Path, package_files: list[Path]) -> None:
    entries = []
    for file_path in sorted(package_files, key=lambda item: item.name):
        content = file_path.read_bytes()
        entries.append(
            {
                "filename": file_path.name,
                "sha256": hashlib.sha256(content).hexdigest(),
                "size_bytes": len(content),
            }
        )
    document = {
        "schema_version": "1.0.0",
        "package_id": "fgrs-pre-ato-system-owner-1.0",
        "system_name": "Federal Grants Review System",
        "system_identifier": "FGRS-PRD",
        "synthetic": True,
        "purpose": "Lowest-common-denominator SSP and control generation evaluation",
        "files": entries,
    }
    path.write_text(json.dumps(document, indent=2) + "\n", encoding="utf-8")


def build_package() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    architecture = OUTPUT_DIR / "diagram-01-production-architecture.png"
    data_flow = OUTPUT_DIR / "diagram-02-business-data-flow.png"
    operations = OUTPUT_DIR / "diagram-03-operations-flow.png"
    _architecture_diagram(architecture)
    _data_flow_diagram(data_flow)
    _operations_diagram(operations)

    product = OUTPUT_DIR / "01-fgrs-product-and-architecture.docx"
    operations_guide = OUTPUT_DIR / "02-fgrs-production-operations-guide.docx"
    access = OUTPUT_DIR / "03-fgrs-user-access-procedures.docx"
    config = OUTPUT_DIR / "04-fgrs-production-config.json"
    inventory = OUTPUT_DIR / "05-fgrs-component-inventory.txt"
    readme = OUTPUT_DIR / "README.md"
    _build_product_overview(product, architecture, data_flow)
    _build_operations_guide(operations_guide, operations)
    _build_access_procedure(access)
    _write_configuration_export(config)
    _write_inventory(inventory)
    _write_readme(readme)

    package_files = [
        product,
        operations_guide,
        access,
        config,
        inventory,
        architecture,
        data_flow,
        operations,
        readme,
    ]
    manifest = OUTPUT_DIR / "manifest.json"
    _write_manifest(manifest, package_files)
    package_files.append(manifest)

    with ZipFile(ZIP_PATH, "w", compression=ZIP_DEFLATED, compresslevel=9) as archive:
        for file_path in sorted(package_files, key=lambda item: item.name):
            archive.write(file_path, arcname=file_path.name)

    print(f"Built {len(package_files)} artifacts in {OUTPUT_DIR}")
    print(f"Archive: {ZIP_PATH}")


if __name__ == "__main__":
    build_package()
