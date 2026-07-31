"""Focused tests for agency template DOCX outline, render, mapping, and review."""

from __future__ import annotations

import asyncio
import json
from io import BytesIO

import pytest
from docx import Document

from ato_service.extraction.types import ExtractionLimits
from ato_service.fisma_generator import GENERIC_DRAFT_NOTICE
from ato_service.ssp_workspace.agency_docx import (
    AgencyDocxError,
    extract_template_outline,
    generate_mapping_plan,
    render_template,
    review_render,
)
from ato_service.ssp_workspace.agency_docx_contracts import (
    ControlTablePlan,
    MappingException,
    MappingPlan,
    TextPlacement,
    canonical_append_column_map,
    control_table_column_names,
)
from ato_service.ssp_workspace.generation import ModelPrompt

LIMITS = ExtractionLimits(
    max_pdf_pages_per_file=10,
    max_extracted_text_characters_per_file=500_000,
    max_zip_members_per_archive=2_000,
    max_zip_uncompressed_bytes_per_archive=50_000_000,
    max_zip_decompression_ratio=100,
    max_xml_depth=128,
    max_xml_elements=500_000,
    max_xml_attributes_per_element=128,
    max_xml_text_node_characters=500_000,
)


def _run(coro: object) -> object:
    return asyncio.run(coro)  # type: ignore[arg-type]


def _approved_snapshot() -> dict[str, object]:
    return {
        "workspace_id": "ws-1",
        "revision_id": "rev-1",
        "content_sha256": "a" * 64,
        "approved_by": "isso@example.gov",
        "approved_at": "2026-01-01T00:00:00Z",
        "system": {
            "display_name": "Grants Portal",
            "external_system_id": "GRANTS-01",
        },
        "profile": {
            "profile_id": "agency-fisma-nist-sp800-53-rev5",
            "version": "1.1.0",
            "impact_level": "moderate",
        },
        "sections": [
            {
                "section_id": "purpose",
                "title": "Purpose",
                "order": 0,
                "state": "approved",
                "content": "The system supports grant processing.",
            }
        ],
        "controls": [
            {
                "control_id": "AC-2",
                "title": "Account Management",
                "state": "approved",
                "implementation_status": "implemented",
                "responsibility": "system_specific",
                "implementation_statement": "Accounts are provisioned centrally.",
                "evidence_links": ["evidence:ac-2"],
            },
            {
                "control_id": "AU-2",
                "title": "Audit Events",
                "state": "approved",
                "implementation_status": "partially_implemented",
                "responsibility": "hybrid",
                "implementation_statement": "Audit events are collected.",
                "evidence_links": [],
            },
        ],
        "questions": [],
    }


def _column_map(**overrides: int) -> dict[str, int]:
    mapping = canonical_append_column_map()
    mapping.update(overrides)
    return mapping


def _build_template(
    *,
    include_control_table: bool = True,
    agency_headers: bool = False,
    table_only: bool = False,
) -> bytes:
    document = Document()
    if not table_only:
        document.add_paragraph("System Name: {{system}}")
        document.add_paragraph("Purpose: {{purpose}}")
    if include_control_table:
        table = document.add_table(rows=2, cols=6)
        if agency_headers:
            headers = [
                "Control Identifier",
                "Control Title",
                "Status",
                "Responsible Party",
                "Implementation Narrative",
                "Evidence",
            ]
        else:
            headers = list(control_table_column_names())
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _mapping_plan(
    *,
    append_table: bool = False,
    column_map: dict[str, int] | None = None,
) -> MappingPlan:
    return MappingPlan(
        text_placements=(
            TextPlacement(
                target_locator="paragraph:0",
                source_ref="system.display_name",
                mode="replace",
            ),
            TextPlacement(
                target_locator="paragraph:1",
                source_ref="section:purpose",
                mode="append",
            ),
        ),
        control_table=ControlTablePlan(
            table_index=None if append_table else 0,
            column_map=column_map or _column_map(),
        ),
        exceptions=(
            MappingException(
                severity="warning",
                code="placeholder",
                message="Purpose placeholder retained.",
            ),
        ),
        summary="Mapped title and purpose.",
    )


def test_extract_template_outline_lists_paragraphs_and_cells() -> None:
    template = _build_template()
    outline = extract_template_outline(template, LIMITS)
    assert outline.paragraphs[0].locator == "paragraph:0"
    assert outline.cells[0].locator == "table:0:cell:0:0"
    assert outline.cells[-1].text == "" or outline.cells[-1].text is not None


def test_render_template_normal_path_preserves_original_and_notice() -> None:
    template = _build_template()
    original = bytes(template)
    rendered = render_template(
        template,
        _mapping_plan(),
        _approved_snapshot(),
        extraction_limits=LIMITS,
    )
    assert template == original
    assert rendered != template
    assert GENERIC_DRAFT_NOTICE in Document(BytesIO(rendered)).paragraphs[0].text
    assert "Grants Portal" in Document(BytesIO(rendered)).paragraphs[1].text


def test_render_template_table_only_inserts_draft_notice() -> None:
    template = _build_template(table_only=True, include_control_table=True)
    plan = MappingPlan(
        text_placements=(),
        control_table=ControlTablePlan(
            table_index=0,
            column_map=_column_map(),
        ),
        exceptions=(),
        summary="Controls only.",
    )
    rendered = render_template(
        template,
        plan,
        _approved_snapshot(),
        extraction_limits=LIMITS,
    )
    document = Document(BytesIO(rendered))
    assert GENERIC_DRAFT_NOTICE in document.paragraphs[0].text


def test_render_template_fills_existing_control_table_with_column_map() -> None:
    template = _build_template(include_control_table=True, agency_headers=True)
    rendered = render_template(
        template,
        _mapping_plan(column_map=_column_map()),
        _approved_snapshot(),
        extraction_limits=LIMITS,
    )
    document = Document(BytesIO(rendered))
    control_table = document.tables[0]
    assert control_table.rows[2].cells[0].text == "AC-2"
    assert control_table.rows[3].cells[0].text == "AU-2"


def test_render_template_appends_control_table_when_index_null() -> None:
    template = _build_template(include_control_table=False)
    rendered = render_template(
        template,
        _mapping_plan(append_table=True),
        _approved_snapshot(),
        extraction_limits=LIMITS,
    )
    document = Document(BytesIO(rendered))
    assert len(document.tables) == 1
    assert document.tables[0].rows[1].cells[0].text == "AC-2"


def test_render_template_rejects_unknown_locator() -> None:
    plan = MappingPlan(
        text_placements=(
            TextPlacement(
                target_locator="paragraph:9",
                source_ref="system.display_name",
                mode="replace",
            ),
        ),
        control_table=ControlTablePlan(
            table_index=None,
            column_map=_column_map(),
        ),
        exceptions=(),
        summary="bad",
    )
    with pytest.raises(AgencyDocxError, match="unknown target locator"):
        render_template(
            _build_template(include_control_table=False),
            plan,
            _approved_snapshot(),
            extraction_limits=LIMITS,
        )


def test_render_template_rejects_duplicate_target() -> None:
    plan = MappingPlan(
        text_placements=(
            TextPlacement(
                target_locator="paragraph:0",
                source_ref="system.display_name",
                mode="replace",
            ),
            TextPlacement(
                target_locator="paragraph:0",
                source_ref="profile.profile_id",
                mode="append",
            ),
        ),
        control_table=ControlTablePlan(
            table_index=None,
            column_map=_column_map(),
        ),
        exceptions=(),
        summary="bad",
    )
    with pytest.raises(AgencyDocxError, match="duplicate target locator"):
        render_template(
            _build_template(include_control_table=False),
            plan,
            _approved_snapshot(),
            extraction_limits=LIMITS,
        )


def test_generate_mapping_plan_uses_one_schema_repair() -> None:
    template = _build_template(include_control_table=False)
    outline = extract_template_outline(template, LIMITS)
    calls: list[ModelPrompt] = []

    def model(prompt: ModelPrompt) -> str:
        calls.append(prompt)
        if len(calls) == 1:
            return json.dumps(
                {
                    "schema_version": "1.0.0",
                    "text_placements": [
                        {
                            "target_locator": "paragraph:0",
                            "source_ref": "system.display_name",
                            "mode": "replace",
                            "extra": "reject-me",
                        }
                    ],
                    "control_table": {
                        "table_index": None,
                        "column_map": _column_map(),
                    },
                    "exceptions": [],
                    "summary": "broken",
                }
            )
        return json.dumps(
            {
                "schema_version": "1.0.0",
                "text_placements": [
                    {
                        "target_locator": "paragraph:0",
                        "source_ref": "system.display_name",
                        "mode": "replace",
                    }
                ],
                "control_table": {
                    "table_index": None,
                    "column_map": _column_map(),
                },
                "exceptions": [],
                "summary": "repaired",
            }
        )

    result = _run(
        generate_mapping_plan(outline, _approved_snapshot(), model)
    )
    assert result.repair_attempted is True
    assert result.attempts == 2
    assert result.plan.summary == "repaired"
    assert len(calls) == 2


def test_generate_mapping_plan_raises_agency_docx_error_on_model_failure() -> None:
    template = _build_template(include_control_table=False)
    outline = extract_template_outline(template, LIMITS)

    def model(_prompt: ModelPrompt) -> str:
        raise RuntimeError("model unavailable")

    with pytest.raises(AgencyDocxError, match="model invocation failed"):
        _run(generate_mapping_plan(outline, _approved_snapshot(), model))


def test_review_render_merges_plan_exceptions_and_model_issues() -> None:
    template = _build_template(include_control_table=False)
    outline = extract_template_outline(template, LIMITS)
    plan = _mapping_plan(append_table=True)
    rendered = render_template(
        template,
        plan,
        _approved_snapshot(),
        extraction_limits=LIMITS,
    )

    def model(_prompt: ModelPrompt) -> str:
        return json.dumps(
            {
                "schema_version": "1.0.0",
                "summary": "Review complete.",
                "issues": [
                    {
                        "severity": "blocker",
                        "code": "missing_mapping",
                        "message": "A required locator stayed empty.",
                        "locator": "paragraph:1",
                    },
                    {
                        "severity": "warning",
                        "code": "draft_notice",
                        "message": "Draft notice present.",
                        "locator": None,
                    },
                ],
            }
        )

    review = _run(
        review_render(outline, plan, _approved_snapshot(), rendered, model)
    )
    assert review.summary == "Review complete."
    assert review.issues[0].code == "placeholder"
    assert review.issues[0].severity == "warning"
    assert any(issue.code == "missing_mapping" for issue in review.issues)
    assert review.facts.section_count == 1
    assert review.facts.control_count == 2
    assert review.facts.plan_exception_count == 1


def test_review_render_adds_blocker_when_control_rows_missing() -> None:
    template = _build_template(include_control_table=False)
    outline = extract_template_outline(template, LIMITS)
    snapshot = _approved_snapshot()
    document = Document(BytesIO(template))
    table = document.add_table(rows=1, cols=6)
    for index, header in enumerate(control_table_column_names()):
        table.rows[0].cells[index].text = header
    buffer = BytesIO()
    document.save(buffer)
    rendered = buffer.getvalue()

    plan = MappingPlan(
        text_placements=(),
        control_table=ControlTablePlan(
            table_index=None,
            column_map=_column_map(),
        ),
        exceptions=(),
        summary="Expected populated control rows.",
    )

    def model(_prompt: ModelPrompt) -> str:
        return json.dumps(
            {
                "schema_version": "1.0.0",
                "summary": "Checked controls.",
                "issues": [],
            }
        )

    review = _run(
        review_render(outline, plan, snapshot, rendered, model)
    )
    assert any(issue.code == "control_table_row_count" for issue in review.issues)
    assert review.issues[0].severity == "blocker"


def test_review_render_raises_agency_docx_error_on_model_failure() -> None:
    template = _build_template(include_control_table=False)
    outline = extract_template_outline(template, LIMITS)
    plan = _mapping_plan(append_table=True)
    rendered = render_template(
        template,
        plan,
        _approved_snapshot(),
        extraction_limits=LIMITS,
    )

    def model(_prompt: ModelPrompt) -> str:
        raise RuntimeError("model unavailable")

    with pytest.raises(AgencyDocxError, match="model invocation failed"):
        _run(
            review_render(
                outline,
                plan,
                _approved_snapshot(),
                rendered,
                model,
            )
        )
