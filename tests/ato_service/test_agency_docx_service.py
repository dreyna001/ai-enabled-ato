"""Service tests for workspace agency DOCX render workflow."""

from __future__ import annotations

import asyncio
import hashlib
import uuid
from datetime import UTC, datetime
from io import BytesIO
from types import SimpleNamespace
from unittest.mock import AsyncMock, MagicMock

import pytest
from docx import Document

from ato_service.blobs import BlobStore
from ato_service.extraction.types import ExtractionLimits
from ato_service.ssp_workspace.agency_docx import AgencyDocxError
from ato_service.ssp_workspace.agency_docx_contracts import (
    ControlTablePlan,
    MappingException,
    MappingPlan,
    ReviewFacts,
    ReviewIssue,
    ReviewResult,
    canonical_append_column_map,
)
from ato_service.ssp_workspace.service import (
    AgencyDocxMalwareScanRequiredError,
    AgencyDocxRenderNotFoundError,
    AgencyDocxRenderStateError,
    AgencyDocxUploadError,
    ApprovalNotFoundError,
    _agency_docx_render_metadata,
    _mapping_plan_document,
    _read_blob_bytes,
    approve_agency_docx_render,
    create_agency_docx_render,
    get_agency_docx_render,
    read_agency_docx_download_bytes,
    read_agency_docx_preview_bytes,
    reject_agency_docx_render,
)

LIMITS = ExtractionLimits(
    max_pdf_pages_per_file=10,
    max_extracted_text_characters_per_file=500_000,
    max_zip_members_per_archive=2_000,
    max_zip_uncompressed_bytes_per_archive=50_000_000,
    max_zip_decompression_ratio=100,
    max_xml_depth=128,
    max_xml_elements=500_000,
    max_xml_attributes_per_element=128,
    max_xml_text_node_characters=500_000,
)


def _run(awaitable: object) -> object:
    return asyncio.run(awaitable)  # type: ignore[arg-type]


def _session() -> MagicMock:
    session = MagicMock()
    session.execute = AsyncMock()
    session.flush = AsyncMock()
    session.add = MagicMock()
    return session


def _template_bytes() -> bytes:
    document = Document()
    document.add_paragraph("System Name:")
    document.add_paragraph("Purpose:")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _snapshot() -> dict[str, object]:
    return {
        "workspace_id": str(uuid.uuid4()),
        "revision_id": str(uuid.uuid4()),
        "content_sha256": "b" * 64,
        "approved_by": "isso@example.gov",
        "approved_at": "2026-01-01T00:00:00Z",
        "system": {
            "display_name": "Grants Portal",
            "external_system_id": "GRANTS-01",
        },
        "profile": {
            "profile_id": "agency-fisma-nist-sp800-53-rev5",
            "version": "1.1.0",
            "impact_level": "moderate",
        },
        "sections": [
            {
                "section_id": "purpose",
                "title": "Purpose",
                "order": 0,
                "state": "approved",
                "content": "The system supports grant processing.",
            }
        ],
        "controls": [
            {
                "control_id": "AC-2",
                "title": "Account Management",
                "state": "approved",
                "implementation_status": "implemented",
                "responsibility": "system_specific",
                "implementation_statement": "Accounts are provisioned centrally.",
                "evidence_links": [],
            }
        ],
        "questions": [],
    }


def _mapping_plan(*, exceptions: tuple = ()) -> MappingPlan:
    return MappingPlan(
        text_placements=(),
        control_table=ControlTablePlan(
            table_index=None,
            column_map=canonical_append_column_map(),
        ),
        exceptions=exceptions,
        summary="Mapped canonical content.",
    )


def _review_result(*, blockers: bool) -> ReviewResult:
    issues: tuple[ReviewIssue, ...]
    if blockers:
        issues = (
            ReviewIssue(
                severity="blocker",
                code="missing_section",
                message="Purpose section not mapped.",
                locator="paragraph:1",
            ),
        )
    else:
        issues = (
            ReviewIssue(
                severity="warning",
                code="sparse_controls",
                message="Only one control mapped.",
                locator=None,
            ),
        )
    return ReviewResult(
        summary="Review complete.",
        issues=issues,
        facts=ReviewFacts(
            section_count=1,
            control_count=1,
            plan_exception_count=0,
            rendered_paragraph_count=3,
            rendered_cell_count=0,
            rendered_table_count=0,
        ),
    )


def _config(*, runtime_profile: str = "dev_local") -> SimpleNamespace:
    return SimpleNamespace(
        limits=SimpleNamespace(max_single_file_bytes=5_000_000),
        document={},
        runtime_profile=runtime_profile,
    )


def test_create_rejects_unapproved_revision(
    tmp_path, monkeypatch: pytest.MonkeyPatch
) -> None:
    session = _session()
    monkeypatch.setattr(
        "ato_service.ssp_workspace.service._approved_export_snapshot",
        AsyncMock(side_effect=ApprovalNotFoundError("approved revision not found")),
    )

    with pytest.raises(ApprovalNotFoundError):
        _run(
            create_agency_docx_render(
                session,
                workspace_id=uuid.uuid4(),
                source_revision_id=uuid.uuid4(),
                template_filename="agency-template.docx",
                template_bytes=_template_bytes(),
                actor_id="isso@example.gov",
                now=datetime.now(UTC),
                blob_store=BlobStore(tmp_path),
                config=_config(),
                model=AsyncMock(),
                audit_hmac_key=b"test-audit-key------",
            )
        )


def test_create_rejects_non_docx_filename(
    tmp_path, monkeypatch: pytest.MonkeyPatch
) -> None:
    session = _session()
    monkeypatch.setattr(
        "ato_service.ssp_workspace.service._approved_export_snapshot",
        AsyncMock(return_value=_snapshot()),
    )

    with pytest.raises(AgencyDocxUploadError):
        _run(
            create_agency_docx_render(
                session,
                workspace_id=uuid.uuid4(),
                source_revision_id=uuid.uuid4(),
                template_filename="agency-template.pdf",
                template_bytes=_template_bytes(),
                actor_id="isso@example.gov",
                now=datetime.now(UTC),
                blob_store=BlobStore(tmp_path),
                config=_config(),
                model=AsyncMock(),
                audit_hmac_key=b"test-audit-key------",
            )
        )


def test_create_returns_exact_cache_hit(tmp_path, monkeypatch: pytest.MonkeyPatch) -> None:
    session = _session()
    workspace_id = uuid.uuid4()
    revision_id = uuid.uuid4()
    cached = SimpleNamespace(render_id=uuid.uuid4())
    workspace = SimpleNamespace(profile_version_id=uuid.uuid4())

    async def execute(statement):  # noqa: ANN001
        sql = str(statement)
        if "ssp_workspaces" in sql:
            result = MagicMock()
            result.scalar_one.return_value = workspace
            return result
        result = MagicMock()
        result.scalar_one_or_none.return_value = cached
        return result

    session.execute = AsyncMock(side_effect=execute)
    monkeypatch.setattr(
        "ato_service.ssp_workspace.service._approved_export_snapshot",
        AsyncMock(return_value=_snapshot()),
    )

    render = _run(
        create_agency_docx_render(
            session,
            workspace_id=workspace_id,
            source_revision_id=revision_id,
            template_filename="agency-template.docx",
            template_bytes=_template_bytes(),
            actor_id="isso@example.gov",
            now=datetime.now(UTC),
            blob_store=BlobStore(tmp_path),
            config=_config(),
            model=AsyncMock(),
            audit_hmac_key=b"test-audit-key------",
        )
    )

    assert render is cached
    session.add.assert_not_called()


def test_create_persists_output_hash_and_review_failed_status(
    tmp_path, monkeypatch: pytest.MonkeyPatch
) -> None:
    session = _session()
    workspace_id = uuid.uuid4()
    revision_id = uuid.uuid4()
    workspace = SimpleNamespace(profile_version_id=uuid.uuid4())
    rendered = _template_bytes()
    output_sha256 = hashlib.sha256(rendered).hexdigest()
    calls = {"count": 0}

    async def execute(statement):  # noqa: ANN001
        sql = str(statement)
        if "ssp_workspaces" in sql:
            result = MagicMock()
            result.scalar_one.return_value = workspace
            return result
        result = MagicMock()
        result.scalar_one_or_none.return_value = None
        return result

    session.execute = AsyncMock(side_effect=execute)

    async def fake_generate(*args, **kwargs):  # noqa: ANN001, ARG001
        return SimpleNamespace(plan=_mapping_plan(), attempts=1, repair_attempted=False)

    async def fake_review(*args, **kwargs):  # noqa: ANN001, ARG001
        return _review_result(blockers=True)

    monkeypatch.setattr(
        "ato_service.ssp_workspace.service._approved_export_snapshot",
        AsyncMock(return_value=_snapshot()),
    )
    monkeypatch.setattr(
        "ato_service.ssp_workspace.agency_docx.generate_mapping_plan",
        fake_generate,
    )
    monkeypatch.setattr(
        "ato_service.ssp_workspace.agency_docx.review_render",
        fake_review,
    )
    monkeypatch.setattr(
        "ato_service.ssp_workspace.agency_docx.render_template",
        lambda *args, **kwargs: rendered,  # noqa: ANN001, ARG001
    )
    monkeypatch.setattr(
        "ato_service.extraction.limits.resolve_extraction_limits_from_config",
        lambda config: LIMITS,  # noqa: ARG001
    )
    monkeypatch.setattr(
        "ato_service.ssp_workspace.service._audit",
        AsyncMock(),
    )

    def capture_add(row: object) -> None:
        calls["count"] += 1
        assert getattr(row, "status") == "review_failed"
        assert getattr(row, "output_sha256") == output_sha256

    session.add = capture_add

    _run(
        create_agency_docx_render(
            session,
            workspace_id=workspace_id,
            source_revision_id=revision_id,
            template_filename="agency-template.docx",
            template_bytes=_template_bytes(),
            actor_id="isso@example.gov",
            now=datetime.now(UTC),
            blob_store=BlobStore(tmp_path),
            config=_config(),
            model=AsyncMock(),
            audit_hmac_key=b"test-audit-key------",
        )
    )

    assert calls["count"] == 1


def test_approve_and_reject_state_transitions(monkeypatch: pytest.MonkeyPatch) -> None:
    session = _session()
    render = SimpleNamespace(
        render_id=uuid.uuid4(),
        status="awaiting_approval",
        mapping_plan={"exceptions": []},
        review_result={"issues": []},
        output_sha256="c" * 64,
        resolved_by=None,
        resolved_at=None,
    )
    monkeypatch.setattr(
        "ato_service.ssp_workspace.service._load_agency_docx_render_for_update",
        AsyncMock(return_value=render),
    )
    monkeypatch.setattr(
        "ato_service.ssp_workspace.service._audit",
        AsyncMock(),
    )

    _run(
        approve_agency_docx_render(
            session,
            workspace_id=uuid.uuid4(),
            render_id=render.render_id,
            actor_id="isso@example.gov",
            now=datetime.now(UTC),
            audit_hmac_key=b"test-audit-key------",
        )
    )
    assert render.status == "approved"
    assert render.resolved_by == "isso@example.gov"

    render.status = "review_failed"
    render.resolved_by = None
    render.resolved_at = None
    _run(
        reject_agency_docx_render(
            session,
            workspace_id=uuid.uuid4(),
            render_id=render.render_id,
            actor_id="isso@example.gov",
            now=datetime.now(UTC),
            audit_hmac_key=b"test-audit-key------",
        )
    )
    assert render.status == "rejected"


def test_approve_rejects_non_pending_render(monkeypatch: pytest.MonkeyPatch) -> None:
    session = _session()
    render = SimpleNamespace(
        render_id=uuid.uuid4(),
        status="review_failed",
        mapping_plan={"exceptions": []},
        review_result={"issues": []},
        output_sha256="c" * 64,
    )
    monkeypatch.setattr(
        "ato_service.ssp_workspace.service._load_agency_docx_render_for_update",
        AsyncMock(return_value=render),
    )

    with pytest.raises(AgencyDocxRenderStateError):
        _run(
            approve_agency_docx_render(
                session,
                workspace_id=uuid.uuid4(),
                render_id=render.render_id,
                actor_id="isso@example.gov",
                now=datetime.now(UTC),
                audit_hmac_key=b"test-audit-key------",
            )
        )


def test_preview_and_download_gates(tmp_path, monkeypatch: pytest.MonkeyPatch) -> None:
    blob_store = BlobStore(tmp_path)
    stored = blob_store.store_stream(BytesIO(_template_bytes()), max_bytes=5_000_000)
    render_id = uuid.uuid4()
    session = _session()
    get_render = AsyncMock(
        side_effect=[
            SimpleNamespace(
                status="awaiting_approval",
                output_storage_key=stored.storage_key,
                output_sha256=stored.sha256,
            ),
            SimpleNamespace(
                status="rejected",
                output_storage_key=stored.storage_key,
                output_sha256=stored.sha256,
            ),
            SimpleNamespace(
                status="rejected",
                output_storage_key=stored.storage_key,
                output_sha256=stored.sha256,
            ),
            SimpleNamespace(
                status="approved",
                output_storage_key=stored.storage_key,
                output_sha256=stored.sha256,
            ),
        ]
    )
    monkeypatch.setattr(
        "ato_service.ssp_workspace.service.get_agency_docx_render",
        get_render,
    )

    preview_bytes = _run(
        read_agency_docx_preview_bytes(
            session,
            workspace_id=uuid.uuid4(),
            render_id=render_id,
            blob_store=blob_store,
        )
    )
    assert hashlib.sha256(preview_bytes).hexdigest() == stored.sha256

    with pytest.raises(AgencyDocxRenderStateError):
        _run(
            read_agency_docx_preview_bytes(
                session,
                workspace_id=uuid.uuid4(),
                render_id=render_id,
                blob_store=blob_store,
            )
        )

    with pytest.raises(AgencyDocxRenderStateError):
        _run(
            read_agency_docx_download_bytes(
                session,
                workspace_id=uuid.uuid4(),
                render_id=render_id,
                blob_store=blob_store,
            )
        )

    download_bytes = _run(
        read_agency_docx_download_bytes(
            session,
            workspace_id=uuid.uuid4(),
            render_id=render_id,
            blob_store=blob_store,
        )
    )
    assert hashlib.sha256(download_bytes).hexdigest() == stored.sha256


def test_get_render_not_found() -> None:
    session = _session()
    result = MagicMock()
    result.scalar_one_or_none.return_value = None
    session.execute = AsyncMock(return_value=result)

    with pytest.raises(AgencyDocxRenderNotFoundError):
        _run(
            get_agency_docx_render(
                session,
                workspace_id=uuid.uuid4(),
                render_id=uuid.uuid4(),
            )
        )


def test_create_translates_mapping_and_review_failures(
    tmp_path, monkeypatch: pytest.MonkeyPatch
) -> None:
    session = _session()
    workspace = SimpleNamespace(profile_version_id=uuid.uuid4())

    async def execute(statement):  # noqa: ANN001
        result = MagicMock()
        if "ssp_workspaces" in str(statement):
            result.scalar_one.return_value = workspace
        else:
            result.scalar_one_or_none.return_value = None
        return result

    session.execute = AsyncMock(side_effect=execute)
    monkeypatch.setattr(
        "ato_service.ssp_workspace.service._approved_export_snapshot",
        AsyncMock(return_value=_snapshot()),
    )
    monkeypatch.setattr(
        "ato_service.extraction.limits.resolve_extraction_limits_from_config",
        lambda config: LIMITS,  # noqa: ARG001
    )
    monkeypatch.setattr(
        "ato_service.ssp_workspace.agency_docx.generate_mapping_plan",
        AsyncMock(side_effect=AgencyDocxError("mapping failed")),
    )

    with pytest.raises(AgencyDocxUploadError, match="mapping failed"):
        _run(
            create_agency_docx_render(
                session,
                workspace_id=uuid.uuid4(),
                source_revision_id=uuid.uuid4(),
                template_filename="agency-template.docx",
                template_bytes=_template_bytes(),
                actor_id="isso@example.gov",
                now=datetime.now(UTC),
                blob_store=BlobStore(tmp_path),
                config=_config(),
                model=AsyncMock(),
                audit_hmac_key=b"test-audit-key------",
            )
        )


def test_create_review_failed_when_mapping_exception_is_blocker(
    tmp_path, monkeypatch: pytest.MonkeyPatch
) -> None:
    session = _session()
    workspace_id = uuid.uuid4()
    revision_id = uuid.uuid4()
    workspace = SimpleNamespace(profile_version_id=uuid.uuid4())
    rendered = _template_bytes()
    captured: dict[str, str] = {}

    async def execute(statement):  # noqa: ANN001
        sql = str(statement)
        if "ssp_workspaces" in sql:
            result = MagicMock()
            result.scalar_one.return_value = workspace
            return result
        result = MagicMock()
        result.scalar_one_or_none.return_value = None
        return result

    session.execute = AsyncMock(side_effect=execute)

    plan = _mapping_plan(
        exceptions=(
            MappingException(
                severity="blocker",
                code="unmapped_section",
                message="Purpose section has no target.",
            ),
        )
    )

    async def fake_generate(*args, **kwargs):  # noqa: ANN001, ARG001
        return SimpleNamespace(plan=plan, attempts=1, repair_attempted=False)

    async def fake_review(*args, **kwargs):  # noqa: ANN001, ARG001
        return _review_result(blockers=False)

    monkeypatch.setattr(
        "ato_service.ssp_workspace.service._approved_export_snapshot",
        AsyncMock(return_value=_snapshot()),
    )
    monkeypatch.setattr(
        "ato_service.ssp_workspace.agency_docx.generate_mapping_plan",
        fake_generate,
    )
    monkeypatch.setattr(
        "ato_service.ssp_workspace.agency_docx.review_render",
        fake_review,
    )
    monkeypatch.setattr(
        "ato_service.ssp_workspace.agency_docx.render_template",
        lambda *args, **kwargs: rendered,  # noqa: ANN001, ARG001
    )
    monkeypatch.setattr(
        "ato_service.extraction.limits.resolve_extraction_limits_from_config",
        lambda config: LIMITS,  # noqa: ARG001
    )
    monkeypatch.setattr(
        "ato_service.ssp_workspace.service._audit",
        AsyncMock(),
    )

    def capture_add(row: object) -> None:
        captured["status"] = getattr(row, "status")

    session.add = capture_add

    _run(
        create_agency_docx_render(
            session,
            workspace_id=workspace_id,
            source_revision_id=revision_id,
            template_filename="agency-template.docx",
            template_bytes=_template_bytes(),
            actor_id="isso@example.gov",
            now=datetime.now(UTC),
            blob_store=BlobStore(tmp_path),
            config=_config(),
            model=AsyncMock(),
            audit_hmac_key=b"test-audit-key------",
        )
    )

    assert captured["status"] == "review_failed"


def test_read_blob_bytes_verifies_payload_digest(tmp_path) -> None:
    from ato_service.storage_reconciliation import require_storage_regular_file

    blob_store = BlobStore(tmp_path)
    stored = blob_store.store_stream(BytesIO(_template_bytes()), max_bytes=5_000_000)
    payload = _read_blob_bytes(blob_store, stored.storage_key, stored.sha256)
    assert hashlib.sha256(payload).hexdigest() == stored.sha256

    prefix, digest = stored.storage_key.split("/", maxsplit=1)
    path = require_storage_regular_file(
        blob_store.storage_root,
        "blobs",
        prefix,
        digest,
    )
    path.write_bytes(b"tampered-bytes")

    with pytest.raises(ValueError, match="digest does not match"):
        _read_blob_bytes(blob_store, stored.storage_key, stored.sha256)


def test_envelope_metadata_exposes_safe_review_and_action_flags() -> None:
    render = SimpleNamespace(
        render_id=uuid.uuid4(),
        profile_version_id=uuid.uuid4(),
        source_revision_id=uuid.uuid4(),
        source_revision_sha256="a" * 64,
        template_sha256="b" * 64,
        template_filename="agency-template.docx",
        output_sha256="c" * 64,
        status="review_failed",
        created_by="isso@example.gov",
        created_at=datetime.now(UTC),
        resolved_by=None,
        resolved_at=None,
        mapping_plan={
            "summary": "Plan summary",
            "exceptions": [
                {
                    "severity": "blocker",
                    "code": "unmapped_section",
                    "message": "Purpose section has no target.",
                }
            ],
            "text_placements": [
                {
                    "target_locator": "paragraph:0",
                    "source_ref": "section:purpose",
                    "mode": "replace",
                }
            ],
        },
        review_result={
            "summary": "Review summary",
            "issues": [
                {
                    "severity": "warning",
                    "code": "sparse_controls",
                    "message": "Only one control mapped.",
                    "locator": None,
                }
            ],
        },
    )

    metadata = _agency_docx_render_metadata(render)

    assert metadata["profile_version_id"] == str(render.profile_version_id)
    assert metadata["can_approve"] is False
    assert metadata["can_preview"] is True
    assert metadata["can_download"] is False
    assert metadata["mapping_exceptions"][0]["severity"] == "blocker"
    assert "text_placements" not in metadata
    assert metadata["review_issues"][0]["severity"] == "warning"


def test_onprem_production_blocks_before_outline_blob_or_model(
    tmp_path, monkeypatch: pytest.MonkeyPatch
) -> None:
    session = _session()
    extract = MagicMock()
    monkeypatch.setattr(
        "ato_service.ssp_workspace.agency_docx.extract_template_outline",
        extract,
    )
    generate = AsyncMock()
    monkeypatch.setattr(
        "ato_service.ssp_workspace.agency_docx.generate_mapping_plan",
        generate,
    )
    blob_store = BlobStore(tmp_path)
    store_calls = {"count": 0}
    original_store = blob_store.store_stream

    def tracked_store(*args: object, **kwargs: object) -> object:
        store_calls["count"] += 1
        return original_store(*args, **kwargs)

    blob_store.store_stream = tracked_store  # type: ignore[method-assign]

    with pytest.raises(AgencyDocxMalwareScanRequiredError):
        _run(
            create_agency_docx_render(
                session,
                workspace_id=uuid.uuid4(),
                source_revision_id=uuid.uuid4(),
                template_filename="agency-template.docx",
                template_bytes=_template_bytes(),
                actor_id="isso@example.gov",
                now=datetime.now(UTC),
                blob_store=blob_store,
                config=_config(runtime_profile="onprem_production"),
                model=AsyncMock(),
                audit_hmac_key=b"test-audit-key------",
            )
        )

    extract.assert_not_called()
    assert store_calls["count"] == 0
    generate.assert_not_called()


def test_reuses_approved_mapping_across_newer_revision(
    tmp_path, monkeypatch: pytest.MonkeyPatch
) -> None:
    session = _session()
    workspace_id = uuid.uuid4()
    new_revision_id = uuid.uuid4()
    profile_version_id = uuid.uuid4()
    workspace = SimpleNamespace(profile_version_id=profile_version_id)
    rendered = _template_bytes()
    reusable_plan = _mapping_plan_document(_mapping_plan())
    reusable = SimpleNamespace(
        mapping_plan=reusable_plan,
        status="approved",
        source_revision_id=uuid.uuid4(),
    )
    new_snapshot = _snapshot()
    new_snapshot["content_sha256"] = "f" * 64
    new_snapshot["revision_id"] = str(new_revision_id)
    agency_queries = {"count": 0}

    async def execute(statement):  # noqa: ANN001
        sql = str(statement)
        if "ssp_workspaces" in sql:
            result = MagicMock()
            result.scalar_one.return_value = workspace
            return result
        if "ssp_agency_docx_renders" in sql:
            agency_queries["count"] += 1
            result = MagicMock()
            if agency_queries["count"] == 1:
                result.scalar_one_or_none.return_value = None
            else:
                result.scalar_one_or_none.return_value = reusable
            return result
        result = MagicMock()
        result.scalar_one_or_none.return_value = None
        return result

    session.execute = AsyncMock(side_effect=execute)
    generate = AsyncMock()
    review = AsyncMock(return_value=_review_result(blockers=False))
    render_mock = MagicMock(return_value=rendered)

    monkeypatch.setattr(
        "ato_service.ssp_workspace.service._approved_export_snapshot",
        AsyncMock(return_value=new_snapshot),
    )
    monkeypatch.setattr(
        "ato_service.ssp_workspace.agency_docx.generate_mapping_plan",
        generate,
    )
    monkeypatch.setattr(
        "ato_service.ssp_workspace.agency_docx.review_render",
        review,
    )
    monkeypatch.setattr(
        "ato_service.ssp_workspace.agency_docx.render_template",
        render_mock,
    )
    monkeypatch.setattr(
        "ato_service.extraction.limits.resolve_extraction_limits_from_config",
        lambda config: LIMITS,  # noqa: ARG001
    )
    monkeypatch.setattr(
        "ato_service.ssp_workspace.service._audit",
        AsyncMock(),
    )

    _run(
        create_agency_docx_render(
            session,
            workspace_id=workspace_id,
            source_revision_id=new_revision_id,
            template_filename="agency-template.docx",
            template_bytes=_template_bytes(),
            actor_id="isso@example.gov",
            now=datetime.now(UTC),
            blob_store=BlobStore(tmp_path),
            config=_config(),
            model=AsyncMock(),
            audit_hmac_key=b"test-audit-key------",
        )
    )

    generate.assert_not_awaited()
    review.assert_awaited_once()
    render_mock.assert_called_once()
    assert render_mock.call_args.args[2] == new_snapshot


def test_approve_rejects_stored_blockers_even_when_status_awaiting(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    session = _session()
    render = SimpleNamespace(
        render_id=uuid.uuid4(),
        status="awaiting_approval",
        mapping_plan={
            "exceptions": [
                {
                    "severity": "blocker",
                    "code": "unmapped_section",
                    "message": "Purpose section has no target.",
                }
            ]
        },
        review_result={"issues": []},
        output_sha256="c" * 64,
        resolved_by=None,
        resolved_at=None,
    )
    monkeypatch.setattr(
        "ato_service.ssp_workspace.service._load_agency_docx_render_for_update",
        AsyncMock(return_value=render),
    )

    with pytest.raises(AgencyDocxRenderStateError, match="blockers"):
        _run(
            approve_agency_docx_render(
                session,
                workspace_id=uuid.uuid4(),
                render_id=render.render_id,
                actor_id="isso@example.gov",
                now=datetime.now(UTC),
                audit_hmac_key=b"test-audit-key------",
            )
        )


def test_can_approve_false_when_blockers_present() -> None:
    metadata = _agency_docx_render_metadata(
        SimpleNamespace(
            render_id=uuid.uuid4(),
            profile_version_id=uuid.uuid4(),
            source_revision_id=uuid.uuid4(),
            source_revision_sha256="a" * 64,
            template_sha256="b" * 64,
            template_filename="agency-template.docx",
            output_sha256="c" * 64,
            status="awaiting_approval",
            created_by="isso@example.gov",
            created_at=datetime.now(UTC),
            resolved_by=None,
            resolved_at=None,
            mapping_plan={
                "summary": "Plan summary",
                "exceptions": [
                    {
                        "severity": "blocker",
                        "code": "unmapped_section",
                        "message": "Purpose section has no target.",
                    }
                ],
            },
            review_result={"summary": "ok", "issues": []},
        )
    )

    assert metadata["can_approve"] is False
