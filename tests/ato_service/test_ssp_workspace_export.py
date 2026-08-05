"""Tests for deterministic SSP workspace JSON and DOCX exports."""

from __future__ import annotations

import asyncio
import hashlib
import json
from io import BytesIO
import uuid
from unittest.mock import AsyncMock

import pytest
from docx import Document

from ato_service.ssp_workspace.export import (
    WorkspaceExportValidationError,
    build_workspace_docx_export,
    build_workspace_json_export,
)
from ato_service.ssp_workspace.sp800_18_docx import nist_control_sort_key
from ato_service.ssp_workspace.oscal_export import (
    OscalSspExportError,
    build_draft_oscal_ssp_json_export,
)
from ato_service.ssp_workspace.service import (
    ApprovalNotFoundError,
    render_approved_export,
)


def _snapshot(*, include_sp800_18: bool = True) -> dict[str, object]:
    snapshot: dict[str, object] = {
        "workspace_id": "11111111-1111-4111-8111-111111111111",
        "revision_id": "22222222-2222-4222-8222-222222222222",
        "content_sha256": "a" * 64,
        "approved_by": "isso@example.gov",
        "approved_at": "2026-07-27T12:00:00Z",
        "document_title": "Grant Intake System",
        "system": {
            "display_name": "Grant Intake System",
            "external_system_id": "GIMS-001",
        },
        "profile": {
            "profile_id": "agency-fisma-80053r5-moderate",
            "version": "5.2.0",
            "impact_level": "moderate",
        },
        "facts": {
            "system.name": "Grant Intake System",
            "system.identifier": "GIMS-001",
            "system.confidentiality_impact": "moderate",
            "system.integrity_impact": "moderate",
            "system.availability_impact": "moderate",
            "system.impact_level": "moderate",
        },
        "sections": [
            {
                "section_id": "system.authorization_boundary",
                "title": "Authorization Boundary",
                "order": 2,
                "state": "reviewed",
                "content": "The boundary includes the application and database.",
            },
            {
                "section_id": "system.purpose",
                "title": "System Purpose",
                "order": 1,
                "state": "reviewed",
                "content": "The system supports grant intake.",
            },
        ],
        "controls": [
            {
                "control_id": "AC-2",
                "title": "Account Management",
                "state": "reviewed",
                "implementation_status": "implemented",
                "responsibility": "hybrid",
                "implementation_statement": "Agency identity manages accounts.",
                "evidence_links": ["artifact-2", "artifact-1", "artifact-1"],
            }
        ],
        "questions": [
            {
                "question_id": "q-2",
                "target": "AU-11",
                "question": "What is the retention period?",
                "owner_type": "agency",
                "status": "open",
            },
            {
                "question_id": "q-1",
                "target": "AC-2",
                "question": "Who reviews accounts?",
                "owner_type": "isso",
                "status": "answered",
            },
        ],
    }
    if include_sp800_18:
        snapshot["standard_coverage"] = [
            {
                "source_id": "nist-sp-800-18r2",
                "requirement_id": "table1.system-name-and-identifier",
                "title": "System Name and Identifier",
                "coverage_kind": "ssp_item",
                "item_ids": ["system.name", "system.identifier"],
                "required": True,
            },
            {
                "source_id": "nist-sp-800-18r2",
                "requirement_id": "table1.system-overview",
                "title": "System Overview",
                "coverage_kind": "ssp_item",
                "item_ids": ["system.purpose"],
                "required": True,
            },
            {
                "source_id": "nist-sp-800-18r2",
                "requirement_id": "table1.authorization-boundary-description",
                "title": "Authorization Boundary Description",
                "coverage_kind": "ssp_item",
                "item_ids": ["system.authorization_boundary"],
                "required": True,
            },
            {
                "source_id": "nist-sp-800-18r2",
                "requirement_id": "table1.system-categorization",
                "title": "System Categorization",
                "coverage_kind": "ssp_item",
                "item_ids": [
                    "system.confidentiality_impact",
                    "system.integrity_impact",
                    "system.availability_impact",
                ],
                "required": True,
            },
            {
                "source_id": "nist-sp-800-18r2",
                "requirement_id": "table1.control-implementation-details",
                "title": "Control Implementation Details",
                "coverage_kind": "controls",
                "item_ids": [],
                "required": True,
            },
            {
                "source_id": "nist-sp-800-18r2",
                "requirement_id": "table1.control-implementation-status",
                "title": "Control Implementation Status",
                "coverage_kind": "controls",
                "item_ids": [],
                "required": True,
            },
        ]
        snapshot["ssp_items"] = [
            {
                "item_id": "system.name",
                "title": "System Name",
                "value_type": "string",
                "standard_refs": ["table1.system-name-and-identifier"],
            },
            {
                "item_id": "system.identifier",
                "title": "System Identifier",
                "value_type": "string",
                "standard_refs": ["table1.system-name-and-identifier"],
            },
            {
                "item_id": "system.purpose",
                "title": "System Purpose",
                "value_type": "string",
                "standard_refs": ["table1.system-overview"],
            },
            {
                "item_id": "system.authorization_boundary",
                "title": "Authorization Boundary",
                "value_type": "string",
                "standard_refs": ["table1.authorization-boundary-description"],
            },
        ]
        snapshot["control_order"] = ["AC-2"]
        snapshot["evidence_catalog"] = {
            "artifact-1": "identity-policy.pdf",
            "artifact-2": "account-review.sarif",
        }
    return snapshot


def test_json_export_is_canonical_and_filters_answered_questions() -> None:
    first = build_workspace_json_export(
        _snapshot(),
        include_open_questions=True,
    )
    second = build_workspace_json_export(
        _snapshot(),
        include_open_questions=True,
    )

    assert hashlib.sha256(first).hexdigest() == hashlib.sha256(second).hexdigest()
    payload = json.loads(first)
    assert [section["section_id"] for section in payload["sections"]] == [
        "system.purpose",
        "system.authorization_boundary",
    ]
    assert payload["schema_version"] == "1.1.0"
    assert payload["controls"][0]["evidence_links"] == [
        "artifact-1",
        "artifact-2",
    ]
    assert [question["question_id"] for question in payload["open_questions"]] == [
        "q-2"
    ]


def test_docx_export_contains_approved_snapshot_content() -> None:
    rendered = build_workspace_docx_export(
        _snapshot(),
        include_open_questions=True,
    )

    document = Document(BytesIO(rendered))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Grant Intake System System Security Plan" in text
    assert "NIST SP 800-18 Revision 2 (Table 1)" in text
    assert "System Overview" in text
    assert "The system supports grant intake." in text
    assert "Authorization Boundary Description" in text
    assert "Control Implementation Details" in text
    assert "AC-2 — Account Management" in text
    assert "identity-policy.pdf" in text
    assert "Control Implementation Status" in text
    assert "What is the retention period?" in text
    assert "Who reviews accounts?" not in text
    assert document.core_properties.author == "isso@example.gov"
    assert len(document.tables) >= 2


def test_docx_export_legacy_flat_layout_without_standard_coverage() -> None:
    rendered = build_workspace_docx_export(
        _snapshot(include_sp800_18=False),
        include_open_questions=False,
    )
    document = Document(BytesIO(rendered))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "System Purpose" in text
    assert "Security Controls" in text


def test_nist_control_sort_key_orders_numeric_segments() -> None:
    assert nist_control_sort_key("AC-3") < nist_control_sort_key("AC-11")
    assert nist_control_sort_key("AC-2(1)") < nist_control_sort_key("AC-2(2)")


def test_export_rejects_duplicate_control_ids() -> None:
    snapshot = _snapshot()
    controls = list(snapshot["controls"])
    controls.append(dict(controls[0]))
    snapshot["controls"] = controls

    with pytest.raises(
        WorkspaceExportValidationError,
        match="duplicate control_id",
    ):
        build_workspace_json_export(snapshot, include_open_questions=False)


def test_export_rejects_invalid_content_hash() -> None:
    snapshot = _snapshot()
    snapshot["content_sha256"] = "not-a-digest"

    with pytest.raises(
        WorkspaceExportValidationError,
        match="content_sha256",
    ):
        build_workspace_docx_export(snapshot, include_open_questions=False)


def test_oscal_json_export_is_schema_valid_draft() -> None:
    payload = json.loads(
        build_draft_oscal_ssp_json_export(
            _snapshot(),
            include_open_questions=True,
        )
    )

    assert "system-security-plan" in payload
    assert payload["system-security-plan"]["metadata"]["oscal-version"]


def test_render_approved_export_requires_approval(monkeypatch: pytest.MonkeyPatch) -> None:
    async def _missing_approval(*_args: object, **_kwargs: object) -> dict[str, object]:
        raise ApprovalNotFoundError("approved revision not found")

    monkeypatch.setattr(
        "ato_service.ssp_workspace.service._approved_export_snapshot",
        _missing_approval,
    )

    with pytest.raises(ApprovalNotFoundError):
        asyncio.run(
            render_approved_export(
                AsyncMock(),
                workspace_id=uuid.uuid4(),
                revision_id=uuid.uuid4(),
                export_format="oscal-json",
                include_open_questions=True,
            )
        )


def test_render_approved_export_delegates_oscal_json_args(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    snapshot = _snapshot()
    delegated: list[tuple[dict[str, object], bool]] = []

    async def _approved_snapshot(*_args: object, **_kwargs: object) -> dict[str, object]:
        return snapshot

    def _oscal_export(
        approved_snapshot: dict[str, object],
        *,
        include_open_questions: bool,
    ) -> bytes:
        delegated.append((approved_snapshot, include_open_questions))
        return b"{}\n"

    monkeypatch.setattr(
        "ato_service.ssp_workspace.service._approved_export_snapshot",
        _approved_snapshot,
    )
    monkeypatch.setattr(
        "ato_service.ssp_workspace.service.build_draft_oscal_ssp_json_export",
        _oscal_export,
    )

    result = asyncio.run(
        render_approved_export(
            AsyncMock(),
            workspace_id=uuid.uuid4(),
            revision_id=uuid.uuid4(),
            export_format="oscal-json",
            include_open_questions=False,
        )
    )

    assert result == b"{}\n"
    assert delegated == [(snapshot, False)]


def test_render_approved_export_rejects_unsupported_format(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    async def _approved_snapshot(*_args: object, **_kwargs: object) -> dict[str, object]:
        return _snapshot()

    monkeypatch.setattr(
        "ato_service.ssp_workspace.service._approved_export_snapshot",
        _approved_snapshot,
    )

    with pytest.raises(ValueError, match="export_format must be json, docx, or oscal-json"):
        asyncio.run(
            render_approved_export(
                AsyncMock(),
                workspace_id=uuid.uuid4(),
                revision_id=uuid.uuid4(),
                export_format="yaml",
                include_open_questions=True,
            )
        )


def test_render_approved_export_surfaces_oscal_validation_without_fallback(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    async def _approved_snapshot(*_args: object, **_kwargs: object) -> dict[str, object]:
        return _snapshot()

    def _oscal_failure(
        _snapshot: dict[str, object],
        *,
        include_open_questions: bool,
    ) -> bytes:
        raise OscalSspExportError("schema validation failed for draft OSCAL export")

    monkeypatch.setattr(
        "ato_service.ssp_workspace.service._approved_export_snapshot",
        _approved_snapshot,
    )
    monkeypatch.setattr(
        "ato_service.ssp_workspace.service.build_draft_oscal_ssp_json_export",
        _oscal_failure,
    )

    with pytest.raises(OscalSspExportError):
        asyncio.run(
            render_approved_export(
                AsyncMock(),
                workspace_id=uuid.uuid4(),
                revision_id=uuid.uuid4(),
                export_format="oscal-json",
                include_open_questions=True,
            )
        )
